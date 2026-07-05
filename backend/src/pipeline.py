"""
Main pipeline for PDF to DOCX reconstruction.
"""

import json
import os
import math
import time
import concurrent.futures
import logging

logging.basicConfig(filename='pipeline_timing.log', level=logging.INFO, format='%(asctime)s - %(message)s')
logger = logging.getLogger(__name__)
from typing import Callable, Dict, List, Optional, Tuple

import numpy as np
import pymupdf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.docx_generator.processors import (
    process_figure_block,
    process_spaced_metadata_row,
    process_table_block,
    process_table_row,
    process_text_block,
    should_merge_with_previous_block,
)
from src.model_loader import ModelConfig, load_doclayout_model
from src.utils import (
    get_section_heading_level,
    is_bbox_contained,
    get_containment_ratio,
    is_same_line,
    horizontally_separated,
)
from src.yolo.iou_matching import LayoutBlock, LayoutRegion, iou, match_blocks_to_layout
from src.yolo.iou_matching import layout_regions_from_detections, text_blocks_from_pdf_elements
from src.yolo.pdf_utils import image_bbox_to_pdf_bbox, render_page_to_image

IMPLICIT_TABLE_EXCLUDED_TYPES = {"plain text", "title"}
METADATA_ROW_TYPES = {"plain text", "title", "section_header"}
PAGE_NUMBER_BLOCK_TYPES = {
    "footer",
    "page_footer",
    "page footer",
    "page_number",
    "page number",
    "page-number",
    "pagination",
}


def _block_text(block: LayoutBlock) -> str:
    text = getattr(block, "text", "") or ""
    if text.strip():
        return " ".join(text.split())

    parts = [
        elem.text.strip()
        for elem in getattr(block, "elements", []) or []
        if getattr(elem, "text", "").strip()
    ]
    return " ".join(" ".join(parts).split())


def _block_line_count(block: LayoutBlock, y_tolerance_pt: float = 2.0) -> int:
    elements = getattr(block, "elements", []) or []
    if not elements:
        text = getattr(block, "text", "") or ""
        return max(1, len([line for line in text.splitlines() if line.strip()]))

    line_tops = []
    for elem in sorted(elements, key=lambda e: (e.bbox[1], e.bbox[0])):
        y0 = elem.bbox[1]
        if not line_tops or abs(y0 - line_tops[-1]) > y_tolerance_pt:
            line_tops.append(y0)

    return max(1, len(line_tops))


def _row_blocks_are_same_line_and_separated(row: List[LayoutBlock]) -> bool:
    if len(row) < 2:
        return False

    for i in range(len(row) - 1):
        if not (
            is_same_line(row[i], row[i + 1])
            and horizontally_separated(row[i], row[i + 1], min_gap=20)
        ):
            return False

    return True


def _is_short_metadata_row(row: List[LayoutBlock]) -> bool:
    if not _row_blocks_are_same_line_and_separated(row):
        return False

    if len(row) > 4:
        return False

    texts = [_block_text(block) for block in row]
    if not all(texts):
        return False

    if not all(block.block_type in METADATA_ROW_TYPES for block in row):
        return False

    if any(len(text) > 120 for text in texts):
        return False

    return any("@" in text for text in texts) or any(
        _block_line_count(block) >= 2 for block in row
    )


def _should_use_implicit_table_row(row: List[LayoutBlock]) -> bool:
    if not _row_blocks_are_same_line_and_separated(row):
        return False

    if any(block.block_type in IMPLICIT_TABLE_EXCLUDED_TYPES for block in row):
        return False

    return True


def _is_original_page_number_block(
    block: LayoutBlock,
    page_width: float,
    page_height: float,
) -> bool:
    text = _block_text(block)
    if not text.isdigit():
        return False

    block_type = (getattr(block, "block_type", "") or "").lower().strip()
    if block_type in PAGE_NUMBER_BLOCK_TYPES:
        return True

    x0, y0, x1, y1 = block.bbox
    block_width = x1 - x0
    block_height = y1 - y0
    block_mid_x = (x0 + x1) / 2.0
    page_mid_x = page_width / 2.0

    centered = abs(block_mid_x - page_mid_x) <= max(18.0, page_width * 0.04)
    near_bottom = y0 >= page_height * 0.88
    footer_sized = block_width <= 40.0 and block_height <= 24.0

    return centered and near_bottom and footer_sized


def _horizontal_overlap_width(bbox_a, bbox_b) -> float:
    return max(0.0, min(bbox_a[2], bbox_b[2]) - max(bbox_a[0], bbox_b[0]))

def _classify_block_column(block: LayoutBlock, page_width: float) -> str:
    """
    Classify a single block as '1col', 'left', or 'right' based on its
    horizontal center's distance from the page center.

    - If the block's center is close to the page center → single-column
    - If the block's center is far to the left → left column
    - If the block's center is far to the right → right column

    The threshold is page_width * 0.15: if the block center is within
    this distance of the page center, it's single-column.
    """
    page_center = page_width / 2.0
    cx = (block.bbox[0] + block.bbox[2]) / 2.0
    deviation = abs(cx - page_center)
    threshold = page_width * 0.25

    if deviation <= threshold:
        return "1col"
    elif cx < page_center:
        return "left"
    else:
        return "right"


def _segment_page_into_bands(layout_blocks: List[LayoutBlock], page_width: float) -> List[Dict]:
    """
    Groups layout blocks into bands of either '1col' or '2col'.

    Detection logic:
      - For each block, compute its horizontal center.
      - If the center is close to the page's horizontal center → single-column.
      - If the center is far left or far right → part of a multi-column section.
      - Consecutive multi-column blocks are grouped into a single '2col' band.
      - Inside a 2col band, blocks are split into left/right columns, then
        paired vertically: each pair shares equivalent vertical starting points.
    """
    if not layout_blocks:
        return []

    sorted_blocks = sorted(layout_blocks, key=lambda b: (b.bbox[1], b.bbox[0]))

    # Step 1: Classify every block
    classified = []  # list of (block, label) where label is '1col', 'left', 'right'
    for b in sorted_blocks:
        label = _classify_block_column(b, page_width)
        classified.append((b, label))

    # Step 2: Group consecutive blocks into runs of same mode
    # A '1col' block always starts/continues a 1col run.
    # A 'left' or 'right' block starts/continues a 2col run.
    bands = []
    current_mode = None  # "1col" or "2col"
    current_1col_blocks = []
    current_left_blocks = []
    current_right_blocks = []

    def _flush_1col():
        nonlocal current_1col_blocks
        if current_1col_blocks:
            bands.append({"mode": "1col", "blocks": list(current_1col_blocks)})
            current_1col_blocks = []

    def _flush_2col():
        nonlocal current_left_blocks, current_right_blocks
        if current_left_blocks or current_right_blocks:
            bands.append({
                "mode": "2col",
                "left_blocks": list(current_left_blocks),
                "right_blocks": list(current_right_blocks),
            })
            current_left_blocks = []
            current_right_blocks = []

    for block, label in classified:
        if label == "1col":
            if current_mode == "2col":
                _flush_2col()
            current_mode = "1col"
            current_1col_blocks.append(block)
        else:  # 'left' or 'right'
            if current_mode == "1col":
                _flush_1col()
            current_mode = "2col"
            if label == "left":
                current_left_blocks.append(block)
            else:
                current_right_blocks.append(block)

    # Flush the last run
    if current_mode == "1col":
        _flush_1col()
    elif current_mode == "2col":
        _flush_2col()

    # Step 3: Post-process 2col bands that have only one side populated.
    # If a band has left blocks but no right blocks (or vice versa), it's
    # really a single-column section that happens to be offset.
    final_bands = []
    for band in bands:
        if band["mode"] == "2col":
            has_left = len(band.get("left_blocks", [])) > 0
            has_right = len(band.get("right_blocks", [])) > 0
            if has_left and has_right:
                final_bands.append(band)
            else:
                # Demote to 1col
                all_blocks = band.get("left_blocks", []) + band.get("right_blocks", [])
                all_blocks.sort(key=lambda b: (b.bbox[1], b.bbox[0]))
                # Merge into adjacent 1col band if possible
                if final_bands and final_bands[-1]["mode"] == "1col":
                    final_bands[-1]["blocks"].extend(all_blocks)
                else:
                    final_bands.append({"mode": "1col", "blocks": all_blocks})
        else:
            # Merge consecutive 1col bands
            if final_bands and final_bands[-1]["mode"] == "1col":
                final_bands[-1]["blocks"].extend(band["blocks"])
            else:
                final_bands.append(band)

    return final_bands


def _trim_table_blocks_against_captions(layout_blocks: List[LayoutBlock]) -> None:
    table_blocks = [block for block in layout_blocks if block.block_type == "table"]
    caption_blocks = [block for block in layout_blocks if block.block_type == "table_caption"]

    for table_block in table_blocks:
        tx0, ty0, tx1, ty1 = table_block.bbox
        table_width = max(1.0, tx1 - tx0)

        for caption_block in caption_blocks:
            cx0, cy0, cx1, _ = caption_block.bbox
            caption_width = max(1.0, cx1 - cx0)
            overlap_width = _horizontal_overlap_width(table_block.bbox, caption_block.bbox)
            enough_horizontal_overlap = overlap_width >= min(table_width, caption_width) * 0.3
            caption_starts_inside_table = ty0 < cy0 < ty1

            if not (caption_starts_inside_table and enough_horizontal_overlap):
                continue

            trimmed_y1 = max(ty0 + 8.0, cy0 - 2.0)
            if trimmed_y1 < ty1:
                table_block.extra["original_yolo_bbox"] = table_block.bbox
                table_block.extra["caption_trimmed"] = True
                table_block.bbox = (tx0, ty0, tx1, trimmed_y1)
                break


class PDFToDocxPipeline:
    """Main pipeline for converting PDF to DOCX."""

    def __init__(
        self,
        model=None,
        style_map: Optional[Dict] = None,
        max_image_width: float = 6.0,
        dpi: int = 300,
        intentional_page_break_gap_threshold_pt: float = 72.0,
    ):
        self.model = model or load_doclayout_model()
        self.style_map = style_map or self._get_default_style_map()
        self.max_image_width = Inches(max_image_width)
        self.dpi = dpi
        self.intentional_page_break_gap_threshold_pt = float(intentional_page_break_gap_threshold_pt)

    @staticmethod
    def _get_default_style_map():
        """Get default style mapping."""
        return {
            "title": {"style": "Title"},
            "section_header": {"style": "Heading 1"},
            "plain text": {"style": "Normal"},
            "abandon": None,
            "figure": None,
            "figure_caption": {"style": "Caption"},
            "table": None,
            "table_caption": {"style": "Caption"},
            "table_footnote": {"style": "Intense Quote"},
            "isolate_formula": {"style": "Normal"},
            "formula_caption": {"style": "Caption"},
        }

    def process_pdf(
        self,
        pdf_path: str,
        output_path: str,
        json_base_path: Optional[str] = None,
        start_page: int = 0,
        end_page: Optional[int] = None,
        progress_callback: Optional[Callable[[int], None]] = None,
    ):
        t_process_pdf_start = time.time()
        pdf_doc = pymupdf.open(pdf_path)
        total_pages = len(pdf_doc)

        if end_page is None:
            end_page = total_pages - 1
        else:
            end_page = min(end_page, total_pages - 1)

        print(f"PDF has {total_pages} pages")
        print(f"Processing pages {start_page} to {end_page} (inclusive)")

        docx_doc = Document()
        self._setup_docx_document(docx_doc)

        section = docx_doc.sections[0]
        try:
            doc_left_margin_in = section.left_margin.inches
        except Exception:
            doc_left_margin_in = section.left_margin.pt / 72.0

        page_width_pts = section.page_width.pt

        first_title_processed = False
        prev_row_y1 = 0
        last_text_paragraph = None
        prev_text_block = None
        prev_page_height = None
        prev_page_last_content_y1 = None
        context = {"last_heading_level": 0}

        for page_idx in range(start_page, end_page):
            progress_callback(35)
            
        print("Rendering pages for YOLO inference...")
        t_render_start = time.time()
        page_images = []
        page_scales = []
        
        def _render_page(p_idx):
            res = render_page_to_image(pdf_doc[p_idx], dpi=self.dpi)
            return p_idx, np.array(res.image), res.scale_x, res.scale_y

        with concurrent.futures.ThreadPoolExecutor(max_workers=os.cpu_count() or 4) as executor:
            futures = [executor.submit(_render_page, i) for i in range(start_page, end_page + 1)]
            results = [f.result() for f in concurrent.futures.as_completed(futures)]
        
        results.sort(key=lambda x: x[0])
        page_images = [r[1] for r in results]
        page_scales = [(r[2], r[3]) for r in results]
        logger.info(f"[Timing] Rendering pages took: {time.time() - t_render_start:.4f}s")
        
        logger.info("Running YOLO inference...")
        t_yolo_start = time.time()
        batch_size = 1
        all_det_dicts = []
        for i in range(0, len(page_images), batch_size):
            batch = page_images[i:i+batch_size]
            yolo_results = self.model.predict(
                batch,
                imgsz=ModelConfig.INPUT_SIZE,
                conf=ModelConfig.CONFIDENCE_THRESHOLD,
                device=ModelConfig.DEVICE,
                verbose=False
            )
            for j, r in enumerate(yolo_results):
                idx = i + j
                scale_x, scale_y = page_scales[idx]
                boxes = r.boxes.xyxy.cpu().numpy()
                scores = r.boxes.conf.cpu().numpy()
                class_ids = r.boxes.cls.cpu().numpy().astype(int)
                names = r.names or self.model.names
                
                det_dicts = []
                for (x0, y0, x1, y1), score, cid in zip(boxes, scores, class_ids):
                    pdf_bbox = image_bbox_to_pdf_bbox((float(x0), float(y0), float(x1), float(y1)), scale_x, scale_y)
                    det_dicts.append({
                        "bbox": pdf_bbox,
                        "score": float(score),
                        "class_id": int(cid),
                        "class_name": str(names.get(int(cid), f"class_{cid}")),
                    })
                all_det_dicts.append(det_dicts)

        logger.info(f"[Timing] YOLO inference (all batches) took: {time.time() - t_yolo_start:.4f}s")

        if progress_callback:
            progress_callback(45)

        t_pages_start = time.time()
        for page_idx in range(start_page, end_page + 1):
            if progress_callback:
                pct = 45 + int(50 * (page_idx - start_page) / max(1, end_page - start_page + 1))
                progress_callback(pct)

            logger.info(f"\nProcessing page {page_idx + 1}/{end_page + 1}...")
            t_page_start = time.time()

            page = pdf_doc[page_idx]
            page_height = page.rect.height

            from src.extract_layout import _get_spans
            t_spans = time.time()
            pdf_elements = _get_spans(page)
            logger.info(f"[Timing] Page {page_idx} _get_spans took: {time.time() - t_spans:.4f}s")
            
            t_images = time.time()
            image_blocks = page.get_image_info()
            for image_index, image_block in enumerate(image_blocks):
                bbox = image_block.get("bbox")
                xref = image_block.get("xref", image_block.get("image"))

                if xref:
                    base_image = pdf_doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    ext = base_image["ext"]
                else:
                    pix = page.get_pixmap(clip=pymupdf.Rect(bbox), dpi=150)
                    image_bytes = pix.tobytes("png")
                    ext = "png"
                
                pdf_elements.append({
                    "type": "image",
                    "bbox": bbox,
                    "xref": xref,
                    "ext": ext,
                    "image_bytes": image_bytes
                })
            logger.info(f"[Timing] Page {page_idx} image extraction took: {time.time() - t_images:.4f}s")

            det_dicts = all_det_dicts[page_idx - start_page]

            result = self._process_page(
                page=page,
                pdf_elements=pdf_elements,
                det_dicts=det_dicts,
                page_idx=page_idx,
                page_folder="",
                docx_doc=docx_doc,
                doc_left_margin_in=doc_left_margin_in,
                page_width_pts=page_width_pts,
                first_title_processed=first_title_processed,
                prev_row_y1=prev_row_y1,
                last_text_paragraph=last_text_paragraph,
                prev_text_block=prev_text_block,
                prev_page_height=prev_page_height,
                prev_page_last_content_y1=prev_page_last_content_y1,
                context=context
            )

            first_title_processed = result["first_title_processed"]
            prev_row_y1 = result["prev_row_y1"]
            last_text_paragraph = result["last_text_paragraph"]
            prev_text_block = result["prev_text_block"]
            prev_page_height = result["prev_page_height"]
            prev_page_last_content_y1 = result["prev_page_last_content_y1"]
            logger.info(f"[Timing] Page {page_idx} _process_page + setup took: {time.time() - t_page_start:.4f}s")

        logger.info(f"[Timing] Total page processing loop took: {time.time() - t_pages_start:.4f}s")
        logger.info("\nAll pages processed!")
        logger.info(f"\nSaving DOCX to: {output_path}")
        t_save_start = time.time()
        docx_doc.save(output_path)
        logger.info(f"[Timing] Saving DOCX took: {time.time() - t_save_start:.4f}s")
        logger.info(f"[Timing] Total process_pdf took: {time.time() - t_process_pdf_start:.4f}s")
        logger.info("Done!")

    def _setup_docx_document(self, docx_doc: Document):
        """Setup DOCX document with footer."""
        section = docx_doc.sections[0]
        footer = section.footer

        if footer.paragraphs:
            footer_paragraph = footer.paragraphs[0]
            for run in list(footer_paragraph.runs):
                footer_paragraph._element.remove(run._element)
        else:
            footer_paragraph = footer.add_paragraph()

        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = footer_paragraph.add_run()

        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'

        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_end)

        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(11)

    def _process_page(
        self,
        page,
        pdf_elements: List[Dict],
        det_dicts: List[Dict],
        page_idx: int,
        page_folder: str,
        docx_doc: Document,
        doc_left_margin_in: float,
        page_width_pts: float,
        first_title_processed: bool,
        prev_row_y1: float,
        last_text_paragraph,
        prev_text_block,
        prev_page_height: Optional[float] = None,
        prev_page_last_content_y1: Optional[float] = None,
        context: Optional[Dict] = None,
    ) -> Dict:
        page_height = page.rect.height

        print(f"  Detected {len(det_dicts)} layout regions")

        text_blocks = text_blocks_from_pdf_elements(pdf_elements)
        layout_regions = layout_regions_from_detections(det_dicts)

        # Extract and match images
        pdf_images = []
        for idx, elem in enumerate(pdf_elements):
            if elem.get("type") == "image":
                bbox = tuple(elem.get("bbox", []))
                if len(bbox) == 4:
                    ext = elem.get("ext", "png") or "png"
                    pdf_images.append({
                        "bbox": bbox,
                        "image_bytes": elem.get("image_bytes"),
                        "index": len(pdf_images),
                        "ext": ext,
                    })

        image_matched_region_indices = set()
        image_to_region_map = {}

        for img_idx, img_data in enumerate(pdf_images):
            img_bbox = img_data["bbox"]
            best_match_idx = None
            best_iou_score = 0.0
            iou_threshold = 0.3

            for region_idx, region in enumerate(layout_regions):
                if region.class_name != "figure":
                    continue
                if region_idx in image_matched_region_indices:
                    continue

                iou_score = iou(img_bbox, region.bbox)
                if iou_score > best_iou_score and iou_score >= iou_threshold:
                    best_iou_score = iou_score
                    best_match_idx = region_idx

            if best_match_idx is not None:
                image_matched_region_indices.add(best_match_idx)
                image_to_region_map[img_idx] = best_match_idx

        layout_blocks = match_blocks_to_layout(
            text_blocks=text_blocks,
            layout_regions=layout_regions,
            iou_threshold=0.1,
        )

        for img_idx, img_data in enumerate(pdf_images):
            if img_idx in image_to_region_map:
                region_idx = image_to_region_map[img_idx]
                matched_region = layout_regions[region_idx]
                layout_block = LayoutBlock(
                    block_type="figure",
                    bbox=img_data["bbox"],
                    text="",
                    score=matched_region.score,
                    elements=[],
                    extra={
                        "image_bytes": img_data.get("image_bytes"),
                        "image_index": img_idx,
                        "matched_region_bbox": matched_region.bbox,
                        "raw_region": matched_region.raw,
                        "matched_blocks": 0,
                    },
                )
            else:
                layout_block = LayoutBlock(
                    block_type="figure",
                    bbox=img_data["bbox"],
                    text="",
                    score=1.0,
                    elements=[],
                    extra={
                        "image_bytes": img_data.get("image_bytes"),
                        "image_index": img_idx,
                        "matched_blocks": 0,
                    },
                )
            layout_blocks.append(layout_block)

        existing_table_region_indices = {
            block.extra.get("layout_region_index")
            for block in layout_blocks
            if block.block_type == "table" and getattr(block, "extra", None)
        }

        for region_idx, region in enumerate(layout_regions):
            if region.class_name == "table":
                if region_idx not in existing_table_region_indices:
                    layout_block = LayoutBlock(
                        block_type="table",
                        bbox=region.bbox,
                        text="",
                        score=region.score,
                        elements=[],
                        extra={
                            "raw_region": region.raw,
                            "layout_region_index": region_idx,
                            "matched_blocks": 0,
                        },
                    )
                    layout_blocks.append(layout_block)

        _trim_table_blocks_against_captions(layout_blocks)

        # Filter out blocks completely contained inside larger image/table blocks
        # This prevents both text and smaller sub-figures from being appended redundantly.
        layout_blocks_filtered = []
        for i, block in enumerate(layout_blocks):
            is_contained = False
            for j, img_block in enumerate(layout_blocks):
                if i == j or img_block.block_type not in ["figure", "table"]:
                    continue
                
                # Check containment using containment ratio > 0.9
                # This handles YOLO bbox noise robustly.
                if get_containment_ratio(block.bbox, img_block.bbox) > 0.90:
                    # If they mutually contain each other (>90% both ways),
                    # only drop the one that appears later in the list to avoid dropping both.
                    if get_containment_ratio(img_block.bbox, block.bbox) > 0.90:
                        if i > j:
                            is_contained = True
                            break
                    else:
                        is_contained = True
                        break

            if not is_contained:
                layout_blocks_filtered.append(block)
        layout_blocks = layout_blocks_filtered

        # Remove empty text blocks
        layout_blocks_cleaned = []
        for block in layout_blocks:
            if block.block_type in ["figure", "table"]:
                layout_blocks_cleaned.append(block)
            elif hasattr(block, "elements") and len(block.elements) > 0:
                layout_blocks_cleaned.append(block)
        layout_blocks = layout_blocks_cleaned

        layout_blocks = [
            block
            for block in layout_blocks
            if not _is_original_page_number_block(block, page.rect.width, page_height)
        ]

        # Margin annotations filter to avoid messing up column detection
        left_strict, right_strict = page.rect.width * 0.05, page.rect.width * 0.95
        left_relaxed, right_relaxed = page.rect.width * 0.10, page.rect.width * 0.90
        filtered_layout_blocks = []
        for b in layout_blocks:
            if b.block_type in ("figure", "table"):
                filtered_layout_blocks.append(b)
                continue
            x0, y0, x1, y1 = b.bbox
            width = x1 - x0
            if x1 < left_strict or x0 > right_strict:
                continue
            if width < 15.0 and (x1 < left_relaxed or x0 > right_relaxed):
                continue
            filtered_layout_blocks.append(b)
        layout_blocks = filtered_layout_blocks

        layout_blocks = sorted(layout_blocks, key=lambda b: (b.bbox[1], b.bbox[0]))

        content_blocks_for_y1 = [b for b in layout_blocks if b.block_type != "abandon"]
        this_page_last_content_y1 = max((b.bbox[3] for b in content_blocks_for_y1), default=0.0)

        for block in layout_blocks:
            if not hasattr(block, 'extra'):
                block.extra = {}
            block.extra['page_idx'] = page_idx
            block.extra['page_height'] = page_height

        bands = _segment_page_into_bands(layout_blocks, page.rect.width)
        print(f"  Segmented into {len(bands)} vertical bands (1col/2col mixed)")

        # Detect intentional page breaks
        first_content_block = next((b for b in layout_blocks if b.block_type != "abandon"), None)
        if (
            first_content_block is not None
            and prev_page_height is not None
            and prev_page_last_content_y1 is not None
        ):
            heading_text = getattr(first_content_block, "text", "") or ""
            starts_section = (
                first_content_block.block_type == "section_header"
                or (first_content_block.block_type == "title" and first_title_processed)
            )
            heading_level, _ = get_section_heading_level(heading_text, default_level=1)
            is_top_level_section = starts_section and (heading_level == 1)

            prev_page_gap_to_bottom = float(prev_page_height) - float(prev_page_last_content_y1)
            prev_page_not_near_bottom = prev_page_gap_to_bottom >= self.intentional_page_break_gap_threshold_pt

            if is_top_level_section and prev_page_not_near_bottom:
                if len(docx_doc.paragraphs) > 0 or len(docx_doc.tables) > 0:
                    docx_doc.add_page_break()

        section = docx_doc.sections[-1]
        from docx.enum.section import WD_SECTION

        for band_idx, band in enumerate(bands):
            mode = band["mode"]
            
            # Switch Word section mode if needed
            current_sect = docx_doc.sections[-1]
            sectPr = current_sect._sectPr
            cols = sectPr.find(qn('w:cols'))
            current_is_2col = (cols is not None and cols.get(qn('w:num')) == '2')

            band_top_y = None
            if mode == "1col":
                valid_blocks = [b for b in band["blocks"] if b.block_type != "abandon"]
                if valid_blocks:
                    band_top_y = valid_blocks[0].bbox[1]
            elif mode == "2col":
                valid_l = [b for b in band.get("left_blocks", []) if b.block_type != "abandon"]
                valid_r = [b for b in band.get("right_blocks", []) if b.block_type != "abandon"]
                l_y = valid_l[0].bbox[1] if valid_l else float('inf')
                r_y = valid_r[0].bbox[1] if valid_r else float('inf')
                if l_y != float('inf') or r_y != float('inf'):
                    band_top_y = min(l_y, r_y)
            
            mode_changed = (mode == "2col" and not current_is_2col) or (mode == "1col" and current_is_2col)
            
            if prev_row_y1 > 0 and band_top_y is not None and mode_changed:
                gap = band_top_y - prev_row_y1
                if gap > 5.0:
                    spacer = docx_doc.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(0)
                    spacer.paragraph_format.space_after = Pt(0)
                    spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    spacer.paragraph_format.line_spacing = Pt(gap)
                    run = spacer.add_run(".")
                    run.font.size = Pt(1)
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    prev_row_y1 = band_top_y

            if mode == "2col" and not current_is_2col:
                new_section = docx_doc.add_section(WD_SECTION.CONTINUOUS)
                new_section.left_margin = Inches(doc_left_margin_in)
                try:
                    new_section.right_margin = current_sect.right_margin
                except:
                    new_section.right_margin = Inches(1.0)
                sectPr = new_section._sectPr
                cols = OxmlElement('w:cols')
                cols.set(qn('w:num'), '2')
                cols.set(qn('w:space'), '360') # 0.25 inch gap
                sectPr.append(cols)
                section = new_section
                last_text_paragraph = None
                prev_text_block = None
            elif mode == "1col" and current_is_2col:
                new_section = docx_doc.add_section(WD_SECTION.CONTINUOUS)
                new_section.left_margin = Inches(doc_left_margin_in)
                try:
                    new_section.right_margin = current_sect.right_margin
                except:
                    new_section.right_margin = Inches(1.0)
                sectPr = new_section._sectPr
                cols = OxmlElement('w:cols')
                cols.set(qn('w:num'), '1')
                sectPr.append(cols)
                section = new_section
                last_text_paragraph = None
                prev_text_block = None

            # Flatten blocks into processing sequence
            blocks_to_process = []
            if mode == "1col":
                for b in band["blocks"]:
                    blocks_to_process.append(b)

                # Group the linear sequence of blocks into rows for processing
                rows = []
                current_row = []
                row_threshold = 10
                for b in blocks_to_process:
                    top_y = b.bbox[1]
                    if not current_row:
                        current_row = [b]
                        current_top = top_y
                    else:
                        if abs(top_y - current_top) <= row_threshold:
                            current_row.append(b)
                        else:
                            rows.append(sorted(current_row, key=lambda x: x.bbox[0]))
                            current_row = [b]
                            current_top = top_y
                if current_row:
                    rows.append(sorted(current_row, key=lambda x: x.bbox[0]))

                for row in rows:
                    row.sort(key=lambda b: b.bbox[0])

                    is_metadata_row = _is_short_metadata_row(row)
                    if is_metadata_row:
                        try:
                            prev_row_y1 = process_spaced_metadata_row(docx_doc, row, section, page_width_pts)
                            last_text_paragraph = None
                            prev_text_block = None
                            continue
                        except Exception:
                            logger.exception("Metadata row processing failed; falling back to text blocks")

                    use_table = _should_use_implicit_table_row(row)

                    if use_table and len(row) >= 2:
                        try:
                            prev_row_y1 = process_table_row(docx_doc, row, section, page_width_pts)
                            continue
                        except Exception:
                            logger.exception("Implicit table row processing failed; falling back to text blocks")

                    for block in row:
                        block_type = block.block_type

                        if block_type == "abandon":
                            continue

                        if block_type == "figure":
                            process_figure_block(
                                docx_doc, block, page,
                                self.max_image_width, page_idx
                            )
                            last_text_paragraph = None
                            prev_text_block = None
                            continue

                        if block_type == "table":
                            process_table_block(
                                docx_doc, block, page,
                                self.max_image_width, page_idx
                            )
                            last_text_paragraph = None
                            prev_text_block = None
                            continue

                        should_merge = False
                        if prev_text_block is not None and last_text_paragraph is not None:
                            should_merge = should_merge_with_previous_block(prev_text_block, block)

                        result = process_text_block(
                            docx_doc, block, block_type, first_title_processed,
                            self.style_map, doc_left_margin_in, row, prev_row_y1,
                            (last_text_paragraph if should_merge else None),
                            context=context
                        )

                        if result[0] is None:
                            continue

                        p, block_type, first_title_processed, y1_pdf, should_merge = result

                        last_text_paragraph = p
                        prev_text_block = block
                        prev_row_y1 = max(prev_row_y1, y1_pdf)

                    valid_bboxes = [b.bbox[3] for b in row if b.block_type != "abandon"]
                    if valid_bboxes:
                        prev_row_y1 = max(valid_bboxes)

            else:
                # 2col band: append left blocks sequentially, then right blocks.
                # Word's 2-column layout will fill the left column first, then
                # flow into the right column automatically.
                from docx.enum.text import WD_BREAK

                left_blocks = sorted(band["left_blocks"], key=lambda b: (b.bbox[1], b.bbox[0]))
                right_blocks = sorted(band["right_blocks"], key=lambda b: (b.bbox[1], b.bbox[0]))

                valid_left = [b for b in left_blocks if b.block_type != "abandon"]
                valid_right = [b for b in right_blocks if b.block_type != "abandon"]
                left_start_y = valid_left[0].bbox[1] if valid_left else float('inf')
                right_start_y = valid_right[0].bbox[1] if valid_right else float('inf')

                left_col_x0 = min((b.bbox[0] for b in valid_left), default=0)
                right_col_x0 = min((b.bbox[0] for b in valid_right), default=0)

                for b in left_blocks:
                    b.extra["in_column"] = True
                    b.extra["col_left_x0_pt"] = left_col_x0
                for b in right_blocks:
                    b.extra["in_column"] = True
                    b.extra["col_left_x0_pt"] = right_col_x0

                # Process left column blocks
                last_text_paragraph = None
                prev_text_block = None
                
                if left_blocks and right_blocks and left_start_y > right_start_y + 5.0:
                    gap = left_start_y - right_start_y
                    spacer = docx_doc.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(gap)
                    spacer.paragraph_format.space_after = Pt(0)
                    spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    spacer.paragraph_format.line_spacing = Pt(1)
                    run = spacer.add_run()
                    run.font.size = Pt(1)

                for block in left_blocks:
                    block_type = block.block_type

                    if block_type == "abandon":
                        continue

                    if block_type == "figure":
                        process_figure_block(
                            docx_doc, block, page,
                            self.max_image_width, page_idx
                        )
                        last_text_paragraph = None
                        prev_text_block = None
                        continue

                    if block_type == "table":
                        process_table_block(
                            docx_doc, block, page,
                            self.max_image_width, page_idx
                        )
                        last_text_paragraph = None
                        prev_text_block = None
                        continue

                    should_merge = False
                    if prev_text_block is not None and last_text_paragraph is not None:
                        should_merge = should_merge_with_previous_block(prev_text_block, block)

                    result = process_text_block(
                        docx_doc, block, block_type, first_title_processed,
                        self.style_map, doc_left_margin_in, [block], prev_row_y1,
                        (last_text_paragraph if should_merge else None),
                        context=context
                    )

                    if result[0] is None:
                        continue

                    p, block_type, first_title_processed, y1_pdf, should_merge = result

                    last_text_paragraph = p
                    prev_text_block = block
                    prev_row_y1 = max(prev_row_y1, y1_pdf)

                # Insert a column break to start the right column.
                # Reset prev_row_y1 to match the vertical starting position of the
                # right column (which should be equivalent to the left column start).
                if right_blocks:
                    p = docx_doc.add_paragraph()
                    run = p.add_run()
                    run.add_break(WD_BREAK.COLUMN)

                    if left_blocks and right_start_y > left_start_y + 5.0:
                        gap = right_start_y - left_start_y
                        spacer = docx_doc.add_paragraph()
                        spacer.paragraph_format.space_before = Pt(gap)
                        spacer.paragraph_format.space_after = Pt(0)
                        spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                        spacer.paragraph_format.line_spacing = Pt(1)
                        run = spacer.add_run(".")
                        run.font.size = Pt(1)
                        from docx.shared import RGBColor
                        run.font.color.rgb = RGBColor(255, 255, 255)

                    # Ensure vertical starting points are equivalent
                    prev_row_y1 = max(left_start_y, right_start_y)
                    last_text_paragraph = None
                    prev_text_block = None

                # Process right column blocks
                for block in right_blocks:
                    block_type = block.block_type

                    if block_type == "abandon":
                        continue

                    if block_type == "figure":
                        process_figure_block(
                            docx_doc, block, page,
                            self.max_image_width, page_idx
                        )
                        last_text_paragraph = None
                        prev_text_block = None
                        continue

                    if block_type == "table":
                        process_table_block(
                            docx_doc, block, page,
                            self.max_image_width, page_idx
                        )
                        last_text_paragraph = None
                        prev_text_block = None
                        continue

                    should_merge = False
                    if prev_text_block is not None and last_text_paragraph is not None:
                        should_merge = should_merge_with_previous_block(prev_text_block, block)

                    result = process_text_block(
                        docx_doc, block, block_type, first_title_processed,
                        self.style_map, doc_left_margin_in, [block], prev_row_y1,
                        (last_text_paragraph if should_merge else None),
                        context=context
                    )

                    if result[0] is None:
                        continue

                    p, block_type, first_title_processed, y1_pdf, should_merge = result

                    last_text_paragraph = p
                    prev_text_block = block
                    prev_row_y1 = max(prev_row_y1, y1_pdf)

                # Update prev_row_y1 to the bottom of the entire 2col band
                all_band_blocks = left_blocks + right_blocks
                valid_band_bboxes = [b.bbox[3] for b in all_band_blocks if b.block_type != "abandon"]
                if valid_band_bboxes:
                    prev_row_y1 = max(valid_band_bboxes)

        return {
            "first_title_processed": first_title_processed,
            "prev_row_y1": prev_row_y1,
            "last_text_paragraph": last_text_paragraph,
            "prev_text_block": prev_text_block,
            "prev_page_height": page_height,
            "prev_page_last_content_y1": this_page_last_content_y1,
        }
