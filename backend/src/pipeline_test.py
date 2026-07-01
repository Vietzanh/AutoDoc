"""
Main pipeline for PDF to DOCX reconstruction.
"""

import json
import os
import math
import time
import concurrent.futures
import logging

logging.basicConfig(filename='pipeline_timing.log', level=logging.INFO, format='%(asctime)s - %(message)s')
logger = logging.getLogger(__name__)
from typing import Callable, Dict, List, Optional, Tuple

import numpy as np
import pymupdf
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.docx_generator.processors import (
    process_figure_block,
    process_spaced_metadata_row,
    process_table_block,
    process_table_row,
    process_text_block,
    should_merge_with_previous_block,
)
from src.model_loader import ModelConfig, load_doclayout_model
from src.utils import (
    get_section_heading_level,
    is_bbox_contained,
    get_containment_ratio,
    is_same_line,
    horizontally_separated,
)
from src.yolo.iou_matching import LayoutBlock, LayoutRegion, iou, match_blocks_to_layout
from src.yolo.iou_matching import layout_regions_from_detections, text_blocks_from_pdf_elements
from src.yolo.pdf_utils import image_bbox_to_pdf_bbox, render_page_to_image

IMPLICIT_TABLE_EXCLUDED_TYPES = {"plain text", "title"}
METADATA_ROW_TYPES = {"plain text", "title", "section_header"}
PAGE_NUMBER_BLOCK_TYPES = {
    "footer",
    "page_footer",
    "page footer",
    "page_number",
    "page number",
    "page-number",
    "pagination",
}

def _detect_2_columns(
    layout_blocks: List[LayoutBlock], page_width: float, strict_overlap: bool = False
) -> Optional[Tuple[List[LayoutBlock], List[LayoutBlock], List[LayoutBlock], List[LayoutBlock]]]:
    """
    Detects a 2-column layout.
    """
    if not layout_blocks or page_width <= 0:
        return None

    page_center = page_width / 2.0
    gutter_width = 15.0  # pt

    wide_breakers: List[LayoutBlock] = []
    tentative_left: List[LayoutBlock] = []
    tentative_right: List[LayoutBlock] = []

    for b in layout_blocks:
        if b.block_type == "abandon":
            continue

        cx = (b.bbox[0] + b.bbox[2]) / 2.0
        block_width = b.bbox[2] - b.bbox[0]

        crosses_gutter = (b.bbox[0] < page_center - gutter_width) and (b.bbox[2] > page_center + gutter_width)
        if crosses_gutter or block_width >= page_width * 0.50:
            wide_breakers.append(b)
        elif cx < page_center:
            tentative_left.append(b)
        else:
            tentative_right.append(b)

    if strict_overlap:
        if not tentative_left or not tentative_right:
            return None
        left_y0 = min(b.bbox[1] for b in tentative_left)
        left_y1 = max(b.bbox[3] for b in tentative_left)
        right_y0 = min(b.bbox[1] for b in tentative_right)
        right_y1 = max(b.bbox[3] for b in tentative_right)
        overlap = min(left_y1, right_y1) - max(left_y0, right_y0)
        if overlap < 10.0:
            return None

    # Pass 2: find true col_zone_top by finding where left and right actually overlap vertically
    overlap_y0s = []
    for lb in tentative_left:
        for rb in tentative_right:
            if lb.bbox[3] > rb.bbox[1] and lb.bbox[1] < rb.bbox[3]:
                overlap_y0s.append(lb.bbox[1])
                overlap_y0s.append(rb.bbox[1])

    col_zone_top = min(overlap_y0s) if overlap_y0s else float('inf')

    final_left: List[LayoutBlock] = []
    final_right: List[LayoutBlock] = []
    positional_breakers: List[LayoutBlock] = []

    for b in tentative_left:
        if b.bbox[3] <= col_zone_top + 15.0:
            positional_breakers.append(b)
        else:
            final_left.append(b)
    for b in tentative_right:
        if b.bbox[3] <= col_zone_top + 15.0:
            positional_breakers.append(b)
        else:
            final_right.append(b)

    if strict_overlap and (not final_left or not final_right):
        return None

    all_breakers = wide_breakers + positional_breakers

    final_left.sort(key=lambda b: (b.bbox[1], b.bbox[0]))
    final_right.sort(key=lambda b: (b.bbox[1], b.bbox[0]))

    if final_left and final_right:
        col_top = min(final_left[0].bbox[1], final_right[0].bbox[1])
    elif final_left:
        col_top = final_left[0].bbox[1]
    elif final_right:
        col_top = final_right[0].bbox[1]
    else:
        col_top = float('inf')

    top_breakers = [b for b in all_breakers if b.bbox[3] <= col_top + 15.0]
    bottom_breakers = [b for b in all_breakers if b not in top_breakers]

    top_breakers.sort(key=lambda b: (b.bbox[1], b.bbox[0]))
    bottom_breakers.sort(key=lambda b: (b.bbox[1], b.bbox[0]))

    return top_breakers, final_left, final_right, bottom_breakers


def _block_text(block: LayoutBlock) -> str:
    text = getattr(block, "text", "") or ""
    if text.strip():
        return " ".join(text.split())

    parts = [
        elem.text.strip()
        for elem in getattr(block, "elements", []) or []
        if getattr(elem, "text", "").strip()
    ]
    return " ".join(" ".join(parts).split())


def _block_line_count(block: LayoutBlock, y_tolerance_pt: float = 2.0) -> int:
    elements = getattr(block, "elements", []) or []
    if not elements:
        text = getattr(block, "text", "") or ""
        return max(1, len([line for line in text.splitlines() if line.strip()]))

    line_tops = []
    for elem in sorted(elements, key=lambda e: (e.bbox[1], e.bbox[0])):
        y0 = elem.bbox[1]
        if not line_tops or abs(y0 - line_tops[-1]) > y_tolerance_pt:
            line_tops.append(y0)

    return max(1, len(line_tops))


def _row_blocks_are_same_line_and_separated(row: List[LayoutBlock]) -> bool:
    if len(row) < 2:
        return False

    for i in range(len(row) - 1):
        if not (
            is_same_line(row[i], row[i + 1])
            and horizontally_separated(row[i], row[i + 1], min_gap=20)
        ):
            return False

    return True


def _is_short_metadata_row(row: List[LayoutBlock]) -> bool:
    if not _row_blocks_are_same_line_and_separated(row):
        return False

    if len(row) > 4:
        return False

    texts = [_block_text(block) for block in row]
    if not all(texts):
        return False

    if not all(block.block_type in METADATA_ROW_TYPES for block in row):
        return False

    if any(len(text) > 120 for text in texts):
        return False

    return any("@" in text for text in texts) or any(
        _block_line_count(block) >= 2 for block in row
    )


def _should_use_implicit_table_row(row: List[LayoutBlock]) -> bool:
    if not _row_blocks_are_same_line_and_separated(row):
        return False

    if any(block.block_type in IMPLICIT_TABLE_EXCLUDED_TYPES for block in row):
        return False

    return True


def _is_original_page_number_block(
    block: LayoutBlock,
    page_width: float,
    page_height: float,
) -> bool:
    text = _block_text(block)
    if not text.isdigit():
        return False

    block_type = (getattr(block, "block_type", "") or "").lower().strip()
    if block_type in PAGE_NUMBER_BLOCK_TYPES:
        return True

    x0, y0, x1, y1 = block.bbox
    block_width = x1 - x0
    block_height = y1 - y0
    block_mid_x = (x0 + x1) / 2.0
    page_mid_x = page_width / 2.0

    centered = abs(block_mid_x - page_mid_x) <= max(18.0, page_width * 0.04)
    near_bottom = y0 >= page_height * 0.88
    footer_sized = block_width <= 40.0 and block_height <= 24.0

    return centered and near_bottom and footer_sized


def _filter_margin_blocks(
    layout_blocks: List[LayoutBlock], page_width: float, page_height: float
) -> List[LayoutBlock]:
    """
    Filter out blocks sitting in page margins (e.g., rotated arXiv sidebar
    annotations like ``arXiv:2606.14730v1 [cs.CV] 2 Jun 2026``).

    Applied *before* column detection so that margin text does not pollute
    the left/right column boundary calculations.
    """
    left_strict = page_width * 0.05
    right_strict = page_width * 0.95
    left_relaxed = page_width * 0.10
    right_relaxed = page_width * 0.90
    narrow_threshold = 15.0  # pt

    filtered: List[LayoutBlock] = []
    for b in layout_blocks:
        # Always keep figure and table blocks
        if b.block_type in ("figure", "table"):
            filtered.append(b)
            continue

        x0, y0, x1, y1 = b.bbox
        block_width = x1 - x0

        # Entire block inside strict left-margin zone
        if x1 < left_strict:
            continue
        # Entire block inside strict right-margin zone
        if x0 > right_strict:
            continue
        # Very narrow block inside relaxed margin zone
        if block_width < narrow_threshold and (x1 < left_relaxed or x0 > right_relaxed):
            continue

        filtered.append(b)

    return filtered


def _horizontal_overlap_width(bbox_a, bbox_b) -> float:
    return max(0.0, min(bbox_a[2], bbox_b[2]) - max(bbox_a[0], bbox_b[0]))


def _trim_table_blocks_against_captions(layout_blocks: List[LayoutBlock]) -> None:
    table_blocks = [block for block in layout_blocks if block.block_type == "table"]
    caption_blocks = [block for block in layout_blocks if block.block_type == "table_caption"]

    for table_block in table_blocks:
        tx0, ty0, tx1, ty1 = table_block.bbox
        table_width = max(1.0, tx1 - tx0)

        for caption_block in caption_blocks:
            cx0, cy0, cx1, _ = caption_block.bbox
            caption_width = max(1.0, cx1 - cx0)
            overlap_width = _horizontal_overlap_width(table_block.bbox, caption_block.bbox)
            enough_horizontal_overlap = overlap_width >= min(table_width, caption_width) * 0.3
            caption_starts_inside_table = ty0 < cy0 < ty1

            if not (caption_starts_inside_table and enough_horizontal_overlap):
                continue

            trimmed_y1 = max(ty0 + 8.0, cy0 - 2.0)
            if trimmed_y1 < ty1:
                table_block.extra["original_yolo_bbox"] = table_block.bbox
                table_block.extra["caption_trimmed"] = True
                table_block.bbox = (tx0, ty0, tx1, trimmed_y1)
                break


class PDFToDocxPipeline:
    """Main pipeline for converting PDF to DOCX."""

    def __init__(
        self,
        model=None,
        style_map: Optional[Dict] = None,
        max_image_width: float = 6.0,
        dpi: int = 300,
        intentional_page_break_gap_threshold_pt: float = 72.0,
    ):
        self.model = model or load_doclayout_model()
        self.style_map = style_map or self._get_default_style_map()
        self.max_image_width = Inches(max_image_width)
        self.dpi = dpi
        self.intentional_page_break_gap_threshold_pt = float(intentional_page_break_gap_threshold_pt)

    @staticmethod
    def _get_default_style_map():
        """Get default style mapping."""
        return {
            "title": {"style": "Title"},
            "section_header": {"style": "Heading 1"},
            "plain text": {"style": "Normal"},
            "abandon": None,
            "figure": None,
            "figure_caption": {"style": "Caption"},
            "table": None,
            "table_caption": {"style": "Caption"},
            "table_footnote": {"style": "Intense Quote"},
            "isolate_formula": {"style": "Normal"},
            "formula_caption": {"style": "Caption"},
        }

    def process_pdf(
        self,
        pdf_path: str,
        output_path: str,
        json_base_path: Optional[str] = None,
        start_page: int = 0,
        end_page: Optional[int] = None,
        progress_callback: Optional[Callable[[int], None]] = None,
    ):
        t_process_pdf_start = time.time()
        pdf_doc = pymupdf.open(pdf_path)
        total_pages = len(pdf_doc)

        if end_page is None:
            end_page = total_pages - 1
        else:
            end_page = min(end_page, total_pages - 1)

        print(f"PDF has {total_pages} pages")
        print(f"Processing pages {start_page} to {end_page} (inclusive)")

        docx_doc = Document()
        self._setup_docx_document(docx_doc)

        # Margins will be configured after the pre-scan

        first_title_processed = False
        prev_row_y1 = 0
        last_text_paragraph = None
        prev_text_block = None
        prev_page_height = None
        prev_page_last_content_y1 = None
        context = {"last_heading_level": 0}

        for page_idx in range(start_page, end_page):
            if progress_callback:
                progress_callback(35)
            
        print("Rendering pages for YOLO inference...")
        t_render_start = time.time()
        page_images = []
        page_scales = []
        
        def _render_page(p_idx):
            res = render_page_to_image(pdf_doc[p_idx], dpi=self.dpi)
            return p_idx, np.array(res.image), res.scale_x, res.scale_y

        with concurrent.futures.ThreadPoolExecutor(max_workers=os.cpu_count() or 4) as executor:
            futures = [executor.submit(_render_page, i) for i in range(start_page, end_page + 1)]
            results = [f.result() for f in concurrent.futures.as_completed(futures)]
        
        results.sort(key=lambda x: x[0])
        page_images = [r[1] for r in results]
        page_scales = [(r[2], r[3]) for r in results]
        logger.info(f"[Timing] Rendering pages took: {time.time() - t_render_start:.4f}s")
        
        logger.info("Running YOLO inference...")
        t_yolo_start = time.time()
        batch_size = 1
        all_det_dicts = []
        for i in range(0, len(page_images), batch_size):
            batch = page_images[i:i+batch_size]
            yolo_results = self.model.predict(
                batch,
                imgsz=ModelConfig.INPUT_SIZE,
                conf=ModelConfig.CONFIDENCE_THRESHOLD,
                device=ModelConfig.DEVICE,
                verbose=False
            )
            for j, r in enumerate(yolo_results):
                idx = i + j
                scale_x, scale_y = page_scales[idx]
                boxes = r.boxes.xyxy.cpu().numpy()
                scores = r.boxes.conf.cpu().numpy()
                class_ids = r.boxes.cls.cpu().numpy().astype(int)
                names = r.names or self.model.names
                
                det_dicts = []
                for (x0, y0, x1, y1), score, cid in zip(boxes, scores, class_ids):
                    pdf_bbox = image_bbox_to_pdf_bbox((float(x0), float(y0), float(x1), float(y1)), scale_x, scale_y)
                    det_dicts.append({
                        "bbox": pdf_bbox,
                        "score": float(score),
                        "class_id": int(cid),
                        "class_name": str(names.get(int(cid), f"class_{cid}")),
                    })
                all_det_dicts.append(det_dicts)

        logger.info(f"[Timing] YOLO inference (all batches) took: {time.time() - t_yolo_start:.4f}s")

        # Global Multi-Column Detection Pre-scan
        self.is_multicolumn_pdf = False
        class DummyBlock:
            def __init__(self, bbox):
                self.bbox = bbox
                self.block_type = "text"
        
        global_min_x0 = float('inf')
        global_max_x1 = 0.0
        left_margins = []
        right_margins = []
        multicol_page_count = 0
        total_pages = len(all_det_dicts)
        
        for i, det_dicts in enumerate(all_det_dicts):
            p_width = pdf_doc[start_page + i].rect.width
            p_height = pdf_doc[start_page + i].rect.height
            dummy_blocks = [DummyBlock(d["bbox"]) for d in det_dicts if d["class_name"] not in ["abandon", "table", "figure"]]
            
            for d in det_dicts:
                x0, _, x1, _ = d["bbox"]
                if 0 <= x0 < p_width:
                    global_min_x0 = min(global_min_x0, x0)
                if 0 < x1 <= p_width:
                    global_max_x1 = max(global_max_x1, x1)
            
            cols_data = _detect_2_columns(dummy_blocks, p_width, strict_overlap=True)
            if cols_data is not None:
                _, left_b, right_b, _ = cols_data
                if len(left_b) >= 5 and len(right_b) >= 5:
                    left_y0 = min(b.bbox[1] for b in left_b)
                    left_y1 = max(b.bbox[3] for b in left_b)
                    right_y0 = min(b.bbox[1] for b in right_b)
                    right_y1 = max(b.bbox[3] for b in right_b)
                    overlap = min(left_y1, right_y1) - max(left_y0, right_y0)
                    
                    if overlap > p_height * 0.2:
                        multicol_page_count += 1
                        left_margins.append(min(b.bbox[0] for b in left_b))
                        right_margins.append(max(b.bbox[2] for b in right_b))
        
        # Require a majority of pages to be 2-column (and at least 2 pages)
        # to declare the whole document as multi-column.
        # This prevents side-by-side author blocks on a single page from
        # fooling the detector.
        if multicol_page_count >= 2 and multicol_page_count > total_pages / 2:
            self.is_multicolumn_pdf = True
                        
        logger.info(f"Global Document Layout: {'Multi-column' if self.is_multicolumn_pdf else 'Single-column'} ({multicol_page_count}/{total_pages} pages detected as 2-col)")
        
        first_page_width = pdf_doc[0].rect.width
        if self.is_multicolumn_pdf and left_margins:
            import statistics
            true_left_x0 = statistics.median(left_margins)
            true_right_x1 = statistics.median(right_margins)
            pdf_left_margin_in = max(0.25, min(true_left_x0 / 72.0, 1.5))
            pdf_right_margin_in = max(0.25, min((first_page_width - true_right_x1) / 72.0, 1.5))
        else:
            pdf_left_margin_in = max(0.25, min(global_min_x0 / 72.0, 1.0))
            pdf_right_margin_in = max(0.25, min((first_page_width - global_max_x1) / 72.0, 1.0))
            
        self.doc_pdf_left_margin_in = pdf_left_margin_in
        self.doc_pdf_right_margin_in = pdf_right_margin_in
        
        # Ensure right margin isn't absurdly small or large
        if getattr(self, "doc_pdf_left_margin_in", None) is None:
            self.doc_pdf_left_margin_in = 1.0
        if getattr(self, "doc_pdf_right_margin_in", None) is None:
            self.doc_pdf_right_margin_in = 1.0

        # Pre-set the document margins only for multi-column PDFs.
        # For single-column, leave the default margins untouched to preserve
        # original behavior (default 1.0 inch margins from template).
        if self.is_multicolumn_pdf:
            for section in docx_doc.sections:
                section.left_margin = Inches(self.doc_pdf_left_margin_in)
                section.right_margin = Inches(self.doc_pdf_right_margin_in)

        section = docx_doc.sections[0]
        if self.is_multicolumn_pdf:
            section.left_margin = Inches(self.doc_pdf_left_margin_in)
            section.right_margin = Inches(self.doc_pdf_right_margin_in)
        try:
            doc_left_margin_in = section.left_margin.inches
        except Exception:
            doc_left_margin_in = section.left_margin.pt / 72.0
            
        page_width_pts = section.page_width.pt

        if progress_callback:
            progress_callback(45)

        t_pages_start = time.time()
        for page_idx in range(start_page, end_page + 1):
            if progress_callback:
                pct = 45 + int(50 * (page_idx - start_page) / max(1, end_page - start_page + 1))
                progress_callback(pct)

            logger.info(f"\nProcessing page {page_idx + 1}/{end_page + 1}...")
            t_page_start = time.time()

            page = pdf_doc[page_idx]
            page_height = page.rect.height

            from src.extract_layout import _get_spans
            t_spans = time.time()
            pdf_elements = _get_spans(page)
            logger.info(f"[Timing] Page {page_idx} _get_spans took: {time.time() - t_spans:.4f}s")
            
            t_images = time.time()
            image_blocks = page.get_image_info()
            for image_index, image_block in enumerate(image_blocks):
                bbox = image_block.get("bbox")
                xref = image_block.get("xref", image_block.get("image"))

                if xref:
                    base_image = pdf_doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    ext = base_image["ext"]
                else:
                    pix = page.get_pixmap(clip=pymupdf.Rect(bbox), dpi=150)
                    image_bytes = pix.tobytes("png")
                    ext = "png"
                
                pdf_elements.append({
                    "type": "image",
                    "bbox": bbox,
                    "xref": xref,
                    "ext": ext,
                    "image_bytes": image_bytes
                })
            logger.info(f"[Timing] Page {page_idx} image extraction took: {time.time() - t_images:.4f}s")

            det_dicts = all_det_dicts[page_idx - start_page]

            result = self._process_page(
                page=page,
                pdf_elements=pdf_elements,
                det_dicts=det_dicts,
                page_idx=page_idx,
                start_page=start_page,
                page_folder="",
                docx_doc=docx_doc,
                doc_left_margin_in=doc_left_margin_in,
                page_width_pts=page_width_pts,
                first_title_processed=first_title_processed,
                prev_row_y1=prev_row_y1,
                last_text_paragraph=last_text_paragraph,
                prev_text_block=prev_text_block,
                prev_page_height=prev_page_height,
                prev_page_last_content_y1=prev_page_last_content_y1,
                context=context,
            )

            first_title_processed = result["first_title_processed"]
            prev_row_y1 = result["prev_row_y1"]
            last_text_paragraph = result["last_text_paragraph"]
            prev_text_block = result["prev_text_block"]
            prev_page_height = result["prev_page_height"]
            prev_page_last_content_y1 = result["prev_page_last_content_y1"]
            logger.info(f"[Timing] Page {page_idx} _process_page + setup took: {time.time() - t_page_start:.4f}s")

        logger.info(f"[Timing] Total page processing loop took: {time.time() - t_pages_start:.4f}s")
        logger.info("\nAll pages processed!")
        logger.info(f"\nSaving DOCX to: {output_path}")
        t_save_start = time.time()
        docx_doc.save(output_path)
        logger.info(f"[Timing] Saving DOCX took: {time.time() - t_save_start:.4f}s")
        logger.info(f"[Timing] Total process_pdf took: {time.time() - t_process_pdf_start:.4f}s")
        logger.info("Done!")

    def _setup_docx_document(self, docx_doc: Document):
        """Setup DOCX document with footer."""
        section = docx_doc.sections[0]
        footer = section.footer

        if footer.paragraphs:
            footer_paragraph = footer.paragraphs[0]
            for run in list(footer_paragraph.runs):
                footer_paragraph._element.remove(run._element)
        else:
            footer_paragraph = footer.add_paragraph()

        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = footer_paragraph.add_run()

        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'

        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_end)

        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(11)

    def _process_page(
        self,
        page,
        pdf_elements: List[Dict],
        det_dicts: List[Dict],
        page_idx: int,
        start_page: int,
        page_folder: str,
        docx_doc: Document,
        doc_left_margin_in: float,
        page_width_pts: float,
        first_title_processed: bool,
        prev_row_y1: float,
        last_text_paragraph,
        prev_text_block,
        context: Dict,
        prev_page_height: Optional[float] = None,
        prev_page_last_content_y1: Optional[float] = None,
    ) -> Dict:
        page_height = page.rect.height

        print(f"  Detected {len(det_dicts)} layout regions")

        text_blocks = text_blocks_from_pdf_elements(pdf_elements)
        layout_regions = layout_regions_from_detections(det_dicts)

        # Extract and match images
        pdf_images = []
        for idx, elem in enumerate(pdf_elements):
            if elem.get("type") == "image":
                bbox = tuple(elem.get("bbox", []))
                if len(bbox) == 4:
                    ext = elem.get("ext", "png") or "png"
                    pdf_images.append({
                        "bbox": bbox,
                        "image_bytes": elem.get("image_bytes"),
                        "index": len(pdf_images),
                        "ext": ext,
                    })

        image_matched_region_indices = set()
        image_to_region_map = {}

        for img_idx, img_data in enumerate(pdf_images):
            img_bbox = img_data["bbox"]
            best_match_idx = None
            best_iou_score = 0.0
            iou_threshold = 0.3

            for region_idx, region in enumerate(layout_regions):
                if region.class_name != "figure":
                    continue
                if region_idx in image_matched_region_indices:
                    continue

                iou_score = iou(img_bbox, region.bbox)
                if iou_score > best_iou_score and iou_score >= iou_threshold:
                    best_iou_score = iou_score
                    best_match_idx = region_idx

            if best_match_idx is not None:
                image_matched_region_indices.add(best_match_idx)
                image_to_region_map[img_idx] = best_match_idx

        layout_blocks = match_blocks_to_layout(
            text_blocks=text_blocks,
            layout_regions=layout_regions,
            iou_threshold=0.1,
        )

        for img_idx, img_data in enumerate(pdf_images):
            if img_idx in image_to_region_map:
                region_idx = image_to_region_map[img_idx]
                matched_region = layout_regions[region_idx]
                layout_block = LayoutBlock(
                    block_type="figure",
                    bbox=img_data["bbox"],
                    text="",
                    score=matched_region.score,
                    elements=[],
                    extra={
                        "image_bytes": img_data.get("image_bytes"),
                        "image_index": img_idx,
                        "matched_region_bbox": matched_region.bbox,
                        "raw_region": matched_region.raw,
                        "matched_blocks": 0,
                    },
                )
            else:
                layout_block = LayoutBlock(
                    block_type="figure",
                    bbox=img_data["bbox"],
                    text="",
                    score=1.0,
                    elements=[],
                    extra={
                        "image_bytes": img_data.get("image_bytes"),
                        "image_index": img_idx,
                        "matched_blocks": 0,
                    },
                )
            layout_blocks.append(layout_block)

        existing_table_region_indices = {
            block.extra.get("layout_region_index")
            for block in layout_blocks
            if block.block_type == "table" and getattr(block, "extra", None)
        }

        for region_idx, region in enumerate(layout_regions):
            if region.class_name == "table":
                if region_idx not in existing_table_region_indices:
                    layout_block = LayoutBlock(
                        block_type="table",
                        bbox=region.bbox,
                        text="",
                        score=region.score,
                        elements=[],
                        extra={
                            "raw_region": region.raw,
                            "layout_region_index": region_idx,
                            "matched_blocks": 0,
                        },
                    )
                    layout_blocks.append(layout_block)

        _trim_table_blocks_against_captions(layout_blocks)

        # Filter out blocks completely contained inside larger image/table blocks
        # This prevents both text and smaller sub-figures from being appended redundantly.
        layout_blocks_filtered = []
        for i, block in enumerate(layout_blocks):
            is_contained = False
            for j, img_block in enumerate(layout_blocks):
                if i == j or img_block.block_type not in ["figure", "table"]:
                    continue
                
                # Check containment using containment ratio > 0.9
                # This handles YOLO bbox noise robustly.
                if get_containment_ratio(block.bbox, img_block.bbox) > 0.90:
                    # If they mutually contain each other (>90% both ways),
                    # only drop the one that appears later in the list to avoid dropping both.
                    if get_containment_ratio(img_block.bbox, block.bbox) > 0.90:
                        if i > j:
                            is_contained = True
                            break
                    else:
                        is_contained = True
                        break

            if not is_contained:
                layout_blocks_filtered.append(block)
        layout_blocks = layout_blocks_filtered

        # Remove empty text blocks
        layout_blocks_cleaned = []
        for block in layout_blocks:
            if block.block_type in ["figure", "table"]:
                layout_blocks_cleaned.append(block)
            elif hasattr(block, "elements") and len(block.elements) > 0:
                layout_blocks_cleaned.append(block)
        layout_blocks = layout_blocks_cleaned

        layout_blocks = [
            block
            for block in layout_blocks
            if not _is_original_page_number_block(block, page.rect.width, page_height)
        ]

        # Filter margin annotations only for multi-column (preserves exact single-col behavior)
        if self.is_multicolumn_pdf:
            layout_blocks = _filter_margin_blocks(layout_blocks, page.rect.width, page_height)

        content_blocks_for_y1 = [b for b in layout_blocks if b.block_type != "abandon"]
        this_page_last_content_y1 = max((b.bbox[3] for b in content_blocks_for_y1), default=0.0)

        for block in layout_blocks:
            if not hasattr(block, 'extra'):
                block.extra = {}
            block.extra['page_idx'] = page_idx
            block.extra['page_height'] = page_height
            block.extra['is_multicolumn_pdf'] = getattr(self, "is_multicolumn_pdf", False)

        if self.is_multicolumn_pdf:
            columns_data = _detect_2_columns(layout_blocks, page.rect.width, strict_overlap=False)
        else:
            columns_data = None

        class DummyBreakBlock:
            def __init__(self, break_type, y):
                self.block_type = break_type
                self.bbox = (0, y, 0, y)
                self.text = ""
                self.score = 1.0
                self.elements = []
                self.extra = {}

        if columns_data is not None:
            top_breakers, left_blocks, right_blocks, bottom_breakers = columns_data
            print("  Detected 2-column layout.")

            # Calculate col_x0 to preserve relative indents
            left_col_x0 = min((b.bbox[0] for b in left_blocks), default=0)
            right_col_x0 = min((b.bbox[0] for b in right_blocks), default=0)

            # Tag all top/bottom breakers as centered to preserve line breaks
            for b in top_breakers + bottom_breakers:
                b.extra["is_centered"] = True

            # Tag column blocks so process_text_block computes relative indent
            for b in left_blocks:
                b.extra["in_column"] = True
                b.extra["col_left_x0_pt"] = left_col_x0
                b.extra["is_multicolumn_pdf"] = True
            for b in right_blocks:
                b.extra["in_column"] = True
                b.extra["col_left_x0_pt"] = right_col_x0
                b.extra["is_multicolumn_pdf"] = True

            rows = []
            col_start_y1 = None
            
            # 1-col section for top breakers (prevents inheriting 2-col from previous page)
            if top_breakers:
                b = DummyBreakBlock("section_1col", 0)
                rows.append([b])
                
                for b_blk in top_breakers: rows.append([b_blk])
                col_start_y1 = max((b_blk.bbox[3] for b_blk in top_breakers), default=0)
            else:
                col_start_y1 = prev_page_last_content_y1 if prev_page_last_content_y1 else 0

            # 2-col section for column content
            b2 = DummyBreakBlock("section_2col", 0)
            rows.append([b2])
            
            for b_blk in left_blocks: rows.append([b_blk])
            
            # column_break with reset_y1 so the right column aligns with the left
            cb_block = DummyBreakBlock("column_break", 0)
            cb_block.extra["reset_y1"] = col_start_y1
            rows.append([cb_block])
            
            for b in right_blocks: rows.append([b])
            
            # 1-col section for bottom breakers
            if bottom_breakers:
                rows.append([DummyBreakBlock("section_1col", 0)])
                for b in bottom_breakers: rows.append([b])
        else:
            layout_blocks = sorted(layout_blocks, key=lambda b: (b.bbox[1], b.bbox[0]))
            
            # Group blocks into rows
            rows = []
            current_row = []
            row_threshold = 10
            for b in layout_blocks:
                top_y = b.bbox[1]
                if not current_row:
                    current_row = [b]
                    current_top = top_y
                else:
                    if abs(top_y - current_top) <= row_threshold:
                        current_row.append(b)
                    else:
                        rows.append(sorted(current_row, key=lambda x: x.bbox[0]))
                        current_row = [b]
                        current_top = top_y
            if current_row:
                rows.append(sorted(current_row, key=lambda x: x.bbox[0]))

        print(f"  Grouped into {len(rows)} rows/items")

        # Detect intentional page breaks
        first_content_block = next((b for b in layout_blocks if b.block_type != "abandon"), None)
        if (
            first_content_block is not None
            and prev_page_height is not None
            and prev_page_last_content_y1 is not None
        ):
            heading_text = getattr(first_content_block, "text", "") or ""
            starts_section = (
                first_content_block.block_type == "section_header"
                or (first_content_block.block_type == "title" and first_title_processed)
            )
            heading_level, _ = get_section_heading_level(heading_text, default_level=1)
            is_top_level_section = starts_section and (heading_level == 1)

            prev_page_gap_to_bottom = float(prev_page_height) - float(prev_page_last_content_y1)
            prev_page_not_near_bottom = prev_page_gap_to_bottom >= self.intentional_page_break_gap_threshold_pt

            if is_top_level_section and prev_page_not_near_bottom:
                if len(docx_doc.paragraphs) > 0 or len(docx_doc.tables) > 0:
                    docx_doc.add_page_break()

        section = docx_doc.sections[0]
        for row in rows:
            row.sort(key=lambda b: b.bbox[0])

            is_metadata_row = _is_short_metadata_row(row)
            if is_metadata_row:
                try:
                    prev_row_y1 = process_spaced_metadata_row(docx_doc, row, section, page_width_pts)
                    last_text_paragraph = None
                    prev_text_block = None
                    continue
                except Exception:
                    logger.exception("Metadata row processing failed; falling back to text blocks")

            use_table = _should_use_implicit_table_row(row)

            if use_table and len(row) >= 2:
                try:
                    prev_row_y1 = process_table_row(docx_doc, row, section, page_width_pts)
                    continue
                except Exception:
                    logger.exception("Implicit table row processing failed; falling back to text blocks")

            for block in row:
                block_type = block.block_type

                if block_type == "section_2col":
                    # Determine if we really need a new section. If the current section is already 2-col, don't add one.
                    # This prevents Word from forcing a page break due to redundant continuous section breaks.
                    current_sect = docx_doc.sections[-1]
                    sectPr = current_sect._sectPr
                    cols = sectPr.find(qn('w:cols'))
                    if cols is not None and cols.get(qn('w:num')) == '2':
                        section = current_sect
                        continue

                    new_section = docx_doc.add_section(WD_SECTION.CONTINUOUS)
                    new_section.left_margin = Inches(self.doc_pdf_left_margin_in)
                    new_section.right_margin = Inches(self.doc_pdf_right_margin_in)
                    sectPr = new_section._sectPr
                    cols = OxmlElement('w:cols')
                    cols.set(qn('w:num'), '2')
                    cols.set(qn('w:space'), '360') # 0.25 inch gap
                    sectPr.append(cols)
                    section = new_section
                    last_text_paragraph = None
                    prev_text_block = None
                    continue

                if block_type == "section_1col":
                    # If current section is already 1-col, don't add another break
                    current_sect = docx_doc.sections[-1]
                    sectPr = current_sect._sectPr
                    cols = sectPr.find(qn('w:cols'))
                    if cols is None or cols.get(qn('w:num')) in [None, '1']:
                        section = current_sect
                        continue

                    new_section = docx_doc.add_section(WD_SECTION.CONTINUOUS)
                    new_section.left_margin = Inches(self.doc_pdf_left_margin_in)
                    new_section.right_margin = Inches(self.doc_pdf_right_margin_in)
                    sectPr = new_section._sectPr
                    cols = OxmlElement('w:cols')
                    cols.set(qn('w:num'), '1')
                    sectPr.append(cols)
                    section = new_section
                    last_text_paragraph = None
                    prev_text_block = None
                    continue

                if block_type == "column_break":
                    if len(docx_doc.paragraphs) > 0:
                        p = docx_doc.paragraphs[-1]
                        if not p.text.strip():
                            run = p.add_run()
                        else:
                            p = docx_doc.add_paragraph()
                            run = p.add_run()
                    else:
                        p = docx_doc.add_paragraph()
                        run = p.add_run()
                    run.add_break(WD_BREAK.COLUMN)
                    
                    if "reset_y1" in block.extra:
                        prev_row_y1 = block.extra["reset_y1"]
                        
                    last_text_paragraph = None
                    prev_text_block = None
                    continue

                if block_type == "abandon":
                    continue

                if block_type == "figure":
                    process_figure_block(
                        docx_doc, block, page,
                        self.max_image_width, page_idx
                    )
                    last_text_paragraph = None
                    prev_text_block = None
                    continue

                if block_type == "table":
                    process_table_block(
                        docx_doc, block, page,
                        self.max_image_width, page_idx
                    )
                    last_text_paragraph = None
                    prev_text_block = None
                    continue

                should_merge = False
                if prev_text_block is not None and last_text_paragraph is not None:
                    should_merge = should_merge_with_previous_block(prev_text_block, block)

                result = process_text_block(
                    docx_doc, block, block_type, first_title_processed,
                    self.style_map, doc_left_margin_in, row, prev_row_y1,
                    (last_text_paragraph if should_merge else None),
                    context=context
                )

                if result[0] is None:
                    continue

                p, block_type, first_title_processed, y1_pdf, should_merge = result

                last_text_paragraph = p
                prev_text_block = block
                prev_row_y1 = max(prev_row_y1, y1_pdf)

            prev_row_y1 = max(b.bbox[3] for b in row)

        return {
            "first_title_processed": first_title_processed,
            "prev_row_y1": prev_row_y1,
            "last_text_paragraph": last_text_paragraph,
            "prev_text_block": prev_text_block,
            "prev_page_height": page_height,
            "prev_page_last_content_y1": this_page_last_content_y1,
        }
