"""
Processors for different block types in DOCX generation.
"""

import os
import io
from typing import Dict, List, Optional, Tuple

import cv2
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor

from src.utils import (
    clean_font_name,
    round_font_size,
    get_section_heading_level,
    remove_table_borders,
    set_table_col_widths,
    is_same_line,
    horizontally_separated,
)
from src.utils.xml_utils import sanitize_text_for_xml
from src.yolo.iou_matching import TextElement, LayoutBlock
import pymupdf


def _block_image_size(block, _max_image_width):
    """Return the block's physical PDF size as DOCX dimensions."""
    x0_pdf, y0_pdf, x1_pdf, y1_pdf = block.bbox
    width_in = max((x1_pdf - x0_pdf) / 72.0, 0.01)
    height_in = max((y1_pdf - y0_pdf) / 72.0, 0.01)

    return Inches(width_in), Inches(height_in)


def _block_has_text(block) -> bool:
    if getattr(block, "text", "") and block.text.strip():
        return True

    for elem in getattr(block, "elements", []) or []:
        if getattr(elem, "text", "") and elem.text.strip():
            return True

    return False


def _remove_table(table) -> None:
    table_element = table._element
    parent = table_element.getparent()
    if parent is not None:
        parent.remove(table_element)


def _line_groups_from_elements(elements, y_tolerance_pt=2.0):
    spans = [
        {"bbox": elem.bbox, "text": elem.text, "element": elem}
        for elem in elements
        if getattr(elem, "text", "").strip()
    ]
    if not spans:
        return []

    spans = sorted(spans, key=lambda s: (s["bbox"][1], s["bbox"][0]))
    lines = []
    current_line = [spans[0]]

    for span in spans[1:]:
        prev = current_line[-1]
        if abs(span["bbox"][1] - prev["bbox"][1]) > y_tolerance_pt:
            lines.append(current_line)
            current_line = [span]
        else:
            current_line.append(span)

    lines.append(current_line)
    return lines


def _metadata_lines_from_block(block):
    if getattr(block, "elements", None):
        return _line_groups_from_elements(block.elements)

    text = getattr(block, "text", "") or ""
    return [
        [
            {
                "bbox": block.bbox,
                "text": line.strip(),
                "element": None,
            }
        ]
        for line in text.splitlines()
        if line.strip()
    ]


def _metadata_line_x0(line, fallback_x0):
    if not line:
        return fallback_x0
    return min(span["bbox"][0] for span in line)


def _rgb_from_pdf_color(color):
    if color is None:
        return RGBColor(0, 0, 0)

    try:
        color_int = int(color)
        if color_int == 16777215:  # 0xFFFFFF (White)
            return RGBColor(0, 0, 0)
    except (TypeError, ValueError):
        return RGBColor(0, 0, 0)

    return RGBColor(
        (color_int >> 16) & 0xFF,
        (color_int >> 8) & 0xFF,
        color_int & 0xFF,
    )


def _apply_span_style(run, elem=None):
    run.font.name = "Times New Roman"
    run.font.color.rgb = _rgb_from_pdf_color(getattr(elem, "color", None))

    if elem is not None and elem.font_size is not None:
        run.font.size = Pt(round_font_size(elem.font_size))

    if elem is not None and elem.font_flags is not None:
        run.bold = (elem.font_flags & 16) != 0
        run.italic = (elem.font_flags & 8) != 0

    return run


def _add_metadata_span_run(paragraph, span):
    run = paragraph.add_run(sanitize_text_for_xml(span["text"]))
    return _apply_span_style(run, span["element"])


def should_merge_with_previous_block(prev_block, curr_block):
    """
    Check if current text block should be merged with previous block into same paragraph.
    """
    if not prev_block or not curr_block:
        return False

    prev_text = prev_block.text if hasattr(prev_block, 'text') else ""
    curr_text = curr_block.text if hasattr(curr_block, 'text') else ""

    if not prev_text or not curr_text:
        return False

    prev_ends_with_space = prev_text.rstrip().endswith(" ")
    prev_no_sentence_end = prev_text.rstrip()[-1] not in ".!?"

    if prev_ends_with_space:
        return True

    curr_starts_lowercase = curr_text.strip() and curr_text.strip()[0].islower()
    return prev_no_sentence_end and curr_starts_lowercase


def process_figure_block(docx_doc, block, pymupdf_page, max_image_width, page_idx):
    """
    Process a figure block by extracting and inserting the image.
    """
    image_path = block.extra.get("image_path") if hasattr(block, "extra") else None
    image_bytes = block.extra.get("image_bytes") if hasattr(block, "extra") else None
    image_width, image_height = _block_image_size(block, max_image_width)

    if image_bytes:
        p = docx_doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(io.BytesIO(image_bytes), width=image_width, height=image_height)
        return True
    elif image_path and os.path.exists(image_path):
        p = docx_doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=image_width, height=image_height)
        return True

    if pymupdf_page is None:
        return False

    x0_pdf, y0_pdf, x1_pdf, y1_pdf = block.bbox
    rect = pymupdf.Rect(x0_pdf, y0_pdf, x1_pdf, y1_pdf)
    
    pix = pymupdf_page.get_pixmap(clip=rect, dpi=150)
    img_bytes = pix.tobytes("png")

    p = docx_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(io.BytesIO(img_bytes), width=image_width, height=image_height)

    return True


def process_table_block(docx_doc, block, pymupdf_page, max_image_width, page_idx):
    """
    Process a table block by cropping and inserting the table as an image.
    """
    if pymupdf_page is None:
        return False

    x0_pdf, y0_pdf, x1_pdf, y1_pdf = block.bbox
    rect = pymupdf.Rect(x0_pdf, y0_pdf, x1_pdf, y1_pdf)
    image_width, image_height = _block_image_size(block, max_image_width)

    pix = pymupdf_page.get_pixmap(clip=rect, dpi=150)
    img_bytes = pix.tobytes("png")

    p = docx_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(io.BytesIO(img_bytes), width=image_width, height=image_height)

    return True


def process_spaced_metadata_row(docx_doc, row, section, page_width_pts):
    """
    Process short same-line metadata blocks as spaced paragraphs, not as a table.
    """
    if not any(_block_has_text(block) for block in row):
        raise ValueError("Metadata row has no extractable text")

    row = sorted(row, key=lambda block: block.bbox[0])
    row_y1 = max(block.bbox[3] for block in row)
    doc_left_margin_pt = section.left_margin.pt
    doc_right_margin_pt = section.right_margin.pt
    available_width_pt = page_width_pts - doc_left_margin_pt - doc_right_margin_pt

    block_lines = [_metadata_lines_from_block(block) for block in row]
    max_lines = max((len(lines) for lines in block_lines), default=0)
    if max_lines == 0:
        raise ValueError("Metadata row produced no DOCX text")

    for line_idx in range(max_lines):
        p = docx_doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

        current_line_tab_positions = []
        for block_idx, block in enumerate(row):
            lines = block_lines[block_idx]
            line = lines[line_idx] if line_idx < len(lines) else []
            line_x0 = _metadata_line_x0(line, block.bbox[0])
            tab_pos_pt = max(1.0, line_x0 - doc_left_margin_pt)
            tab_pos_pt = min(tab_pos_pt, max(1.0, available_width_pt - 1.0))
            current_line_tab_positions.append(tab_pos_pt)

        for tab_pos_pt in current_line_tab_positions:
            p.paragraph_format.tab_stops.add_tab_stop(Pt(tab_pos_pt))

        for block_idx, lines in enumerate(block_lines):
            p.add_run("\t")

            if line_idx >= len(lines):
                continue

            for span in lines[line_idx]:
                _add_metadata_span_run(p, span)

    return row_y1


def process_table_row(docx_doc, row, section, page_width_pts):
    """
    Process a row of blocks as a table.
    """
    if not any(_block_has_text(block) for block in row):
        raise ValueError("Implicit table row has no extractable text")

    min_x_pt = min(b.bbox[0] for b in row)
    max_x_pt = max(b.bbox[2] for b in row)
    doc_left_margin_pt = section.left_margin.pt
    doc_right_margin_pt = section.right_margin.pt

    table_start_x = doc_left_margin_pt
    table_end_x = page_width_pts - doc_right_margin_pt
    table_total_width_pt = table_end_x - table_start_x

    num_cols = len(row)
    table = docx_doc.add_table(rows=1, cols=num_cols)
    table.autofit = False
    remove_table_borders(table)

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.append(tblPr)

    for e in tblPr.findall(qn("w:tblInd")):
        tblPr.remove(e)

    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), "0")
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)

    block_widths = []
    block_positions = []
    for b in row:
        block_w = max(1.0, (b.bbox[2] - b.bbox[0]))
        block_widths.append(block_w)
        block_positions.append((b.bbox[0], b.bbox[2]))

    first_block_start = block_positions[0][0]
    last_block_end = block_positions[-1][1]
    total_block_span = last_block_end - first_block_start

    col_widths_pt = []
    if total_block_span > 0 and len(block_widths) > 0:
        for i, block_w in enumerate(block_widths):
            proportion = block_w / total_block_span
            col_width = proportion * table_total_width_pt
            col_widths_pt.append(max(1.0, col_width))

        current_total = sum(col_widths_pt)
        if current_total < table_total_width_pt:
            remaining = table_total_width_pt - current_total
            for i in range(len(col_widths_pt)):
                proportion = block_widths[i] / sum(block_widths) if sum(block_widths) > 0 else 1.0 / len(block_widths)
                col_widths_pt[i] += proportion * remaining
    else:
        col_width_per_col = table_total_width_pt / num_cols
        col_widths_pt = [col_width_per_col] * num_cols

    set_table_col_widths(table, col_widths_pt)

    inserted_text = False

    for col_idx, block in enumerate(row):
        cell = table.rows[0].cells[col_idx]
        cell.text = ""

        if not block.elements:
            fallback_text = getattr(block, "text", "") or ""
            if fallback_text.strip():
                for line in fallback_text.splitlines():
                    if not line.strip():
                        continue
                    p = cell.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(sanitize_text_for_xml(line.strip()))
                    _apply_span_style(run)
                    inserted_text = True
            continue

        elems = block.elements
        spans = [{"bbox": e.bbox, "text": e.text, "element": e} for e in elems]
        spans = sorted(spans, key=lambda s: (s["bbox"][1], s["bbox"][0]))

        lines = []
        current_line = [spans[0]]
        y_tolerance_pt = 2.0

        for s in spans[1:]:
            prev = current_line[-1]
            y_prev = prev["bbox"][1]
            y_curr = s["bbox"][1]
            if abs(y_curr - y_prev) > y_tolerance_pt:
                lines.append(current_line)
                current_line = [s]
            else:
                current_line.append(s)
        lines.append(current_line)

        for line in lines:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for span in line:
                elem = span["element"]
                text_content = span["text"]
                if not text_content.endswith(" "):
                    text_content += " "

                run = p.add_run(sanitize_text_for_xml(text_content))
                if text_content.strip():
                    inserted_text = True
                _apply_span_style(run, elem)

    if not inserted_text:
        _remove_table(table)
        raise ValueError("Implicit table row produced no DOCX text")

    return max(b.bbox[3] for b in row)


def process_text_block(
    docx_doc, block, block_type, first_title_processed,
    style_map, doc_left_margin_in, row, prev_row_y1, last_paragraph=None,
    context=None
):
    """
    Process a text block (title, plain text, captions, etc.).
    """
    if block_type == "title":
        if not first_title_processed:
            first_title_processed = True
        else:
            block_type = "section_header"

    style_info = style_map.get(block_type, {"style": "Normal"})
    if style_info is None:
        return None, block_type, first_title_processed, block.bbox[3], False

    x0_pdf, y0_pdf, x1_pdf, y1_pdf = block.bbox

    section = docx_doc.sections[0]
    try:
        doc_right_margin_in = section.right_margin.inches
    except Exception:
        doc_right_margin_in = section.right_margin.pt / 72.0
    page_width_in = section.page_width.pt / 72.0

    x0_in = x0_pdf / 72
    x1_in = x1_pdf / 72
    indent_from_margin_in = max(0.0, x0_in - doc_left_margin_in - (5.0 / 72.0))
    # Re-enable right indent with a small tolerance buffer (0.1 inch extra)
    # to account for DOCX font rendering requiring slightly more width than PDF.
    right_edge_in = page_width_in - doc_right_margin_in
    raw_right_indent = right_edge_in - x1_in - (5.0 / 72.0)
    right_indent_from_margin_in = max(0.0, raw_right_indent - 0.1)

    # --- Justify detection ---
    # Group elements into visual lines, then check if most lines share the
    # same right-edge x-coordinate (within tolerance). If so, the PDF author
    # used justified alignment.
    is_justified = False
    if block.elements and block_type not in ["title", "figure_caption", "table_caption", "formula_caption"]:
        lines: Dict[float, List] = {}
        line_height = 0.0
        for elem in block.elements:
            eh = elem.bbox[3] - elem.bbox[1]
            if eh > line_height:
                line_height = eh
        half_h = max(line_height * 0.5, 3.0)
        for elem in block.elements:
            cy = (elem.bbox[1] + elem.bbox[3]) / 2.0
            matched = False
            for key in lines:
                if abs(cy - key) < half_h:
                    lines[key].append(elem)
                    matched = True
                    break
            if not matched:
                lines[cy] = [elem]
        if len(lines) >= 3:
            # For each line, find the rightmost span end
            right_edges = []
            for key in sorted(lines.keys()):
                line_elems = lines[key]
                max_x1 = max(e.bbox[2] for e in line_elems)
                right_edges.append(max_x1)
            # Exclude the last line (typically shorter in justified text)
            edges_to_check = right_edges[:-1]
            if len(edges_to_check) >= 2:
                # Check if most lines end at the same x (within 5pt tolerance)
                ref_edge = edges_to_check[0]
                same_count = sum(1 for e in edges_to_check if abs(e - ref_edge) < 5.0)
                if same_count / len(edges_to_check) >= 0.6:
                    is_justified = True

    should_merge = (last_paragraph is not None)

    style_name = style_info["style"]

    if should_merge and last_paragraph is not None:
        p = last_paragraph
    else:
        if style_name.startswith("Heading"):
            base_level = int(style_name.split()[-1])
            if block_type == "section_header":
                heading_text = getattr(block, "text", "") or ""
                heading_level, is_numbered = get_section_heading_level(heading_text, default_level=base_level)
                
                if context is not None:
                    if is_numbered:
                        context["last_heading_level"] = heading_level
                    else:
                        last_level = context.get("last_heading_level", 0)
                        # If unnumbered, make it a sub-header of the last numbered section
                        heading_level = last_level + 1 if last_level > 0 else 1
                        # We don't update last_heading_level for unnumbered sub-headers
                        # so subsequent unnumbered headers stay at the same level
            else:
                heading_level = base_level
            p = docx_doc.add_heading(level=heading_level)
        else:
            p = docx_doc.add_paragraph(style=style_name)

        p.paragraph_format.left_indent = Inches(indent_from_margin_in)
        p.paragraph_format.right_indent = Inches(right_indent_from_margin_in)

        if block == row[0] and prev_row_y1 > 0:
            vertical_gap = max(0, (y0_pdf - prev_row_y1))
            p.paragraph_format.space_before = Pt(vertical_gap)
        else:
            p.paragraph_format.space_before = Pt(0)

        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        if block_type == "title":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif block_type in ["figure_caption", "table_caption"]:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif block_type not in ["formula_caption"] and not style_name.startswith("Heading"):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if is_justified else WD_ALIGN_PARAGRAPH.LEFT

    if not block.elements:
        return p, block_type, first_title_processed, y1_pdf, should_merge

    span_to_block_idx = block.extra.get("span_to_block_idx")

    if span_to_block_idx is not None:
        text_block_groups: Dict[int, List[Tuple[TextElement, int]]] = {}

        span_to_original_order = block.extra.get("span_to_original_order")

        for elem_idx, elem in enumerate(block.elements):
            block_idx = span_to_block_idx.get(elem_idx, 0)
            if block_idx not in text_block_groups:
                text_block_groups[block_idx] = []

            if span_to_original_order and elem_idx in span_to_original_order:
                _, original_span_idx = span_to_original_order[elem_idx]
            else:
                original_span_idx = elem_idx

            text_block_groups[block_idx].append((elem, original_span_idx))

        sorted_block_indices = sorted(text_block_groups.keys())

        for group_idx, block_idx in enumerate(sorted_block_indices):
            group_with_order = sorted(text_block_groups[block_idx], key=lambda x: x[1])
            group = [elem for elem, _ in group_with_order]

            if group_idx > 0:
                p = docx_doc.add_paragraph(style=style_name)
                p.paragraph_format.left_indent = Inches(indent_from_margin_in)
                p.paragraph_format.right_indent = Inches(right_indent_from_margin_in)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                if block_type == "title":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif block_type in ["figure_caption", "table_caption"]:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif block_type not in ["formula_caption"] and not style_name.startswith("Heading"):
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if is_justified else WD_ALIGN_PARAGRAPH.LEFT

            prev_elem = None
            for elem_i, elem in enumerate(group):
                if prev_elem is not None:
                    is_new_line = elem.bbox[1] > prev_elem.bbox[1] + (prev_elem.bbox[3] - prev_elem.bbox[1]) * 0.5
                    if is_new_line:
                        prev_ended_early = (x1_pdf - prev_elem.bbox[2]) > 15.0
                        curr_indented = (elem.bbox[0] - x0_pdf) > 15.0
                        if prev_ended_early or curr_indented:
                            p.add_run().add_break()
                
                text_content = elem.text
                # Only add trailing space if the next span is NOT horizontally
                # adjacent (i.e. part of the same word split by font-size change).
                needs_space = True
                if not text_content.endswith(" ") and elem_i + 1 < len(group):
                    next_elem = group[elem_i + 1]
                    next_on_same_line = not (next_elem.bbox[1] > elem.bbox[1] + (elem.bbox[3] - elem.bbox[1]) * 0.5)
                    if next_on_same_line:
                        gap = next_elem.bbox[0] - elem.bbox[2]
                        font_size = elem.font_size or 12.0
                        # If gap is less than ~40% of font size, spans are part of the same word
                        if gap < font_size * 0.4:
                            needs_space = False
                if needs_space and not text_content.endswith(" "):
                    text_content += " "
                run = p.add_run(sanitize_text_for_xml(text_content))
                _apply_span_style(run, elem)
                prev_elem = elem
    else:
        font_sizes = [e.font_size for e in block.elements if e.font_size is not None]
        avg_font_size = sum(font_sizes) / len(font_sizes) if font_sizes else 12.0
        vertical_gap_threshold = avg_font_size * 1.5

        text_block_groups = []
        current_group = [block.elements[0]]

        for i in range(1, len(block.elements)):
            prev_elem = block.elements[i - 1]
            curr_elem = block.elements[i]

            vertical_gap = curr_elem.bbox[1] - prev_elem.bbox[3]

            if vertical_gap > vertical_gap_threshold:
                text_block_groups.append(current_group)
                current_group = [curr_elem]
            else:
                current_group.append(curr_elem)

        if current_group:
            text_block_groups.append(current_group)

        for group_idx, group in enumerate(text_block_groups):
            if group_idx > 0:
                p = docx_doc.add_paragraph(style=style_name)
                p.paragraph_format.left_indent = Inches(indent_from_margin_in)
                p.paragraph_format.right_indent = Inches(right_indent_from_margin_in)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                if block_type == "title":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif block_type in ["figure_caption", "table_caption"]:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif block_type not in ["formula_caption"] and not style_name.startswith("Heading"):
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if is_justified else WD_ALIGN_PARAGRAPH.LEFT

            prev_elem = None
            for elem_i, elem in enumerate(group):
                if prev_elem is not None:
                    is_new_line = elem.bbox[1] > prev_elem.bbox[1] + (prev_elem.bbox[3] - prev_elem.bbox[1]) * 0.5
                    if is_new_line:
                        prev_ended_early = (x1_pdf - prev_elem.bbox[2]) > 15.0
                        curr_indented = (elem.bbox[0] - x0_pdf) > 15.0
                        if prev_ended_early or curr_indented:
                            p.add_run().add_break()

                text_content = elem.text
                # Only add trailing space if the next span is NOT horizontally
                # adjacent (i.e. part of the same word split by font-size change).
                needs_space = True
                if not text_content.endswith(" ") and elem_i + 1 < len(group):
                    next_elem = group[elem_i + 1]
                    next_on_same_line = not (next_elem.bbox[1] > elem.bbox[1] + (elem.bbox[3] - elem.bbox[1]) * 0.5)
                    if next_on_same_line:
                        gap = next_elem.bbox[0] - elem.bbox[2]
                        font_size = elem.font_size or 12.0
                        if gap < font_size * 0.4:
                            needs_space = False
                if needs_space and not text_content.endswith(" "):
                    text_content += " "
                run = p.add_run(sanitize_text_for_xml(text_content))
                _apply_span_style(run, elem)
                prev_elem = elem

    return p, block_type, first_title_processed, y1_pdf, should_merge
