"""
Main pipeline for PDF to DOCX reconstruction.
"""

import json
import os
from typing import Dict, List, Optional, Tuple

import numpy as np
import pymupdf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from src.docx_generator.processors import (
        process_figure_block,
        process_table_block,
        process_table_row,
        process_text_block,
        should_merge_with_previous_block,
    )
    from src.model_loader import ModelConfig, load_doclayout_model
    from src.utils import (
        get_section_heading_level,
        is_bbox_contained,
        is_same_line,
        horizontally_separated,
    )
    from src.yolo.iou_matching import LayoutBlock, LayoutRegion, iou, match_blocks_to_layout
    from src.yolo.iou_matching import layout_regions_from_detections, text_blocks_from_pdf_elements
    from src.yolo.pdf_utils import image_bbox_to_pdf_bbox, render_page_to_image
except ImportError:
    # Fallback for direct execution
    import sys
    from pathlib import Path
    sys.path.insert(0, str(Path(__file__).parent))
    from docx_generator.processors import (
        process_figure_block,
        process_table_block,
        process_table_row,
        process_text_block,
        should_merge_with_previous_block,
    )
    from model_loader import ModelConfig, load_doclayout_model
    from utils import (
        get_section_heading_level,
        is_bbox_contained,
        is_same_line,
        horizontally_separated,
    )
    from yolo.iou_matching import LayoutBlock, LayoutRegion, iou, match_blocks_to_layout
    from yolo.iou_matching import layout_regions_from_detections, text_blocks_from_pdf_elements
    from yolo.pdf_utils import image_bbox_to_pdf_bbox, render_page_to_image


class PDFToDocxPipeline:
    """Main pipeline for converting PDF to DOCX."""
    
    def __init__(
        self,
        model=None,
        style_map: Optional[Dict] = None,
        max_image_width: float = 6.0,
        dpi: int = 300,
        intentional_page_break_gap_threshold_pt: float = 72.0,
    ):
        """
        Initialize pipeline.
        
        Args:
            model: YOLOv10 model instance (if None, will be loaded)
            style_map: Dictionary mapping block types to DOCX styles
            max_image_width: Maximum image width in inches
            dpi: DPI for rendering PDF pages
        """
        self.model = model or load_doclayout_model()
        self.style_map = style_map or self._get_default_style_map()
        self.max_image_width = Inches(max_image_width)
        self.dpi = dpi
        # If the previous page's last content ends at least this many points above
        # the bottom of the page, and the next page starts a new top-level section,
        # we treat that as an intentional author page break.
        self.intentional_page_break_gap_threshold_pt = float(intentional_page_break_gap_threshold_pt)
        
    @staticmethod
    def _get_default_style_map():
        """Get default style mapping."""
        return {
            "title": {"style": "Title"},
            "section_header": {"style": "Heading 1"},
            "plain text": {"style": "Normal"},
            "abandon": None,  # Ignore
            "figure": None,  # Handled separately as image
            "figure_caption": {"style": "Caption"},
            "table": None,  # Handled separately as table
            "table_caption": {"style": "Caption"},
            "table_footnote": {"style": "Intense Quote"},
            "isolate_formula": {"style": "Normal"},
            "formula_caption": {"style": "Caption"},
        }
    
    def process_pdf(
        self,
        pdf_path: str,
        output_path: str,
        json_base_path: Optional[str] = None,
        start_page: int = 0,
        end_page: Optional[int] = None,
    ):
        """
        Process PDF file and generate DOCX.
        
        Args:
            pdf_path: Path to input PDF file
            output_path: Path to output DOCX file
            json_base_path: Base path for JSON metadata files (defaults to data_layout)
            start_page: First page to process (0-indexed)
            end_page: Last page to process (None = all pages)
        """
        # Open PDF
        pdf_doc = pymupdf.open(pdf_path)
        total_pages = len(pdf_doc)
        
        # Determine page range
        if end_page is None:
            end_page = total_pages - 1
        else:
            end_page = min(end_page, total_pages - 1)
        
        print(f"PDF has {total_pages} pages")
        print(f"Processing pages {start_page} to {end_page} (inclusive)")
        
        # Create DOCX document
        docx_doc = Document()
        self._setup_docx_document(docx_doc)
        
        # Get document settings
        section = docx_doc.sections[0]
        try:
            doc_left_margin_in = section.left_margin.inches
        except Exception:
            doc_left_margin_in = section.left_margin.pt / 72.0
        
        page_width_pts = section.page_width.pt
        
        # Track state across pages
        first_title_processed = False
        prev_row_y1 = 0
        last_text_paragraph = None
        prev_text_block = None
        prev_page_height = None
        prev_page_last_content_y1 = None
        
        # Process each page
        for page_idx in range(start_page, end_page + 1):
            print(f"\nProcessing page {page_idx + 1}/{end_page + 1}...")
            
            # Load JSON metadata
            if json_base_path is None:
                json_base_path = "../data_layout"
            page_folder = os.path.join(json_base_path, f"page_{page_idx}")
            json_path = os.path.join(page_folder, f"page_{page_idx}_layout.json")
            
            if not os.path.exists(json_path):
                print(f"Warning: JSON file not found for page {page_idx}, skipping...")
                continue
            
            with open(json_path, "r", encoding="utf-8") as f:
                pdf_elements = json.load(f)
            
            # Get page from PDF
            page = pdf_doc[page_idx]
            page_height = page.rect.height
            
            # Process page
            result = self._process_page(
                page=page,
                pdf_elements=pdf_elements,
                page_idx=page_idx,
                page_folder=page_folder,
                docx_doc=docx_doc,
                doc_left_margin_in=doc_left_margin_in,
                page_width_pts=page_width_pts,
                first_title_processed=first_title_processed,
                prev_row_y1=prev_row_y1,
                last_text_paragraph=last_text_paragraph,
                prev_text_block=prev_text_block,
                prev_page_height=prev_page_height,
                prev_page_last_content_y1=prev_page_last_content_y1,
            )
            
            # Update state
            first_title_processed = result["first_title_processed"]
            prev_row_y1 = result["prev_row_y1"]
            last_text_paragraph = result["last_text_paragraph"]
            prev_text_block = result["prev_text_block"]
            prev_page_height = result["prev_page_height"]
            prev_page_last_content_y1 = result["prev_page_last_content_y1"]
        
        print("\nAll pages processed!")
        
        # Save DOCX
        print(f"\nSaving DOCX to: {output_path}")
        docx_doc.save(output_path)
        print("Done!")
    
    def _setup_docx_document(self, docx_doc: Document):
        """Setup DOCX document with footer."""
        section = docx_doc.sections[0]
        footer = section.footer
        
        # Get or create footer paragraph
        if footer.paragraphs:
            footer_paragraph = footer.paragraphs[0]
            # Clear existing content
            for run in list(footer_paragraph.runs):
                footer_paragraph._element.remove(run._element)
        else:
            footer_paragraph = footer.add_paragraph()
        
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page number field
        run = footer_paragraph.add_run()
        
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_end)
        
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(11)
    
    def _process_page(
        self,
        page,
        pdf_elements: List[Dict],
        page_idx: int,
        page_folder: str,
        docx_doc: Document,
        doc_left_margin_in: float,
        page_width_pts: float,
        first_title_processed: bool,
        prev_row_y1: float,
        last_text_paragraph,
        prev_text_block,
        prev_page_height: Optional[float] = None,
        prev_page_last_content_y1: Optional[float] = None,
    ) -> Dict:
        """
        Process a single page.
        
        Returns:
            Dict with updated state: first_title_processed, prev_row_y1, last_text_paragraph, prev_text_block
        """
        page_height = page.rect.height
        
        # Render page to image
        render_result = render_page_to_image(page, dpi=self.dpi)
        page_image_pil = render_result.image
        scale_x, scale_y = render_result.scale_x, render_result.scale_y
        page_image = np.array(page_image_pil)
        
        # Run YOLO detection
        results = self.model.predict(
            page_image,
            imgsz=ModelConfig.INPUT_SIZE,
            conf=ModelConfig.CONFIDENCE_THRESHOLD,
            device=ModelConfig.DEVICE,
        )
        
        r = results[0]
        boxes = r.boxes.xyxy.cpu().numpy()
        scores = r.boxes.conf.cpu().numpy()
        class_ids = r.boxes.cls.cpu().numpy().astype(int)
        names = r.names or self.model.names
        
        # Convert to detection dict format
        det_dicts = []
        for (x0, y0, x1, y1), score, cid in zip(boxes, scores, class_ids):
            det_dicts.append({
                "bbox": (float(x0), float(y0), float(x1), float(y1)),
                "score": float(score),
                "class_id": int(cid),
                "class_name": str(names.get(int(cid), f"class_{cid}")),
            })
        
        # Convert YOLO bboxes from image space to PDF coordinate space
        for det in det_dicts:
            x0, y0, x1, y1 = det["bbox"]
            pdf_bbox = image_bbox_to_pdf_bbox((x0, y0, x1, y1), scale_x, scale_y)
            det["bbox"] = pdf_bbox
        
        print(f"  Detected {len(det_dicts)} layout regions")
        
        # Build text blocks and layout regions
        text_blocks = text_blocks_from_pdf_elements(pdf_elements)
        layout_regions = layout_regions_from_detections(det_dicts)
        
        # Extract and match images
        pdf_images = []
        for idx, elem in enumerate(pdf_elements):
            if elem.get("type") == "image":
                bbox = tuple(elem.get("bbox", []))
                if len(bbox) == 4:
                    ext = elem.get("ext", "png") or "png"
                    image_path = os.path.join(page_folder, "images", f"img_{len(pdf_images)}.{ext}")
                    pdf_images.append({
                        "bbox": bbox,
                        "path": image_path,
                        "index": len(pdf_images),
                        "ext": ext,
                    })
        
        # Match images with figure regions
        image_matched_region_indices = set()
        image_to_region_map = {}
        
        for img_idx, img_data in enumerate(pdf_images):
            img_bbox = img_data["bbox"]
            best_match_idx = None
            best_iou_score = 0.0
            iou_threshold = 0.3
            
            for region_idx, region in enumerate(layout_regions):
                if region.class_name != "figure":
                    continue
                if region_idx in image_matched_region_indices:
                    continue
                
                iou_score = iou(img_bbox, region.bbox)
                if iou_score > best_iou_score and iou_score >= iou_threshold:
                    best_iou_score = iou_score
                    best_match_idx = region_idx
            
            if best_match_idx is not None:
                image_matched_region_indices.add(best_match_idx)
                image_to_region_map[img_idx] = best_match_idx
        
        # Match text blocks with layout regions
        layout_blocks = match_blocks_to_layout(
            text_blocks=text_blocks,
            layout_regions=layout_regions,
            iou_threshold=0.1,
        )
        
        # Create LayoutBlocks for images
        for img_idx, img_data in enumerate(pdf_images):
            if img_idx in image_to_region_map:
                region_idx = image_to_region_map[img_idx]
                matched_region = layout_regions[region_idx]
                layout_block = LayoutBlock(
                    block_type="figure",
                    bbox=matched_region.bbox,
                    text="",
                    score=matched_region.score,
                    elements=[],
                    extra={
                        "image_path": img_data["path"],
                        "image_index": img_idx,
                        "raw_region": matched_region.raw,
                        "matched_blocks": 0,
                    },
                )
            else:
                layout_block = LayoutBlock(
                    block_type="figure",
                    bbox=img_data["bbox"],
                    text="",
                    score=1.0,
                    elements=[],
                    extra={
                        "image_path": img_data["path"],
                        "image_index": img_idx,
                        "matched_blocks": 0,
                    },
                )
            layout_blocks.append(layout_block)
        
        # Create LayoutBlocks for table regions
        for region in layout_regions:
            if region.class_name == "table":
                already_exists = False
                for existing_block in layout_blocks:
                    if existing_block.block_type == "table":
                        iou_score = iou(existing_block.bbox, region.bbox)
                        if iou_score > 0.9:
                            already_exists = True
                            break
                
                if not already_exists:
                    layout_block = LayoutBlock(
                        block_type="table",
                        bbox=region.bbox,
                        text="",
                        score=region.score,
                        elements=[],
                        extra={"raw_region": region.raw, "matched_blocks": 0},
                    )
                    layout_blocks.append(layout_block)
        
        # Filter text blocks contained in image/table blocks
        image_blocks = [b for b in layout_blocks if b.block_type == "figure"]
        table_blocks = [b for b in layout_blocks if b.block_type == "table"]
        layout_blocks_filtered = []
        for block in layout_blocks:
            if block.block_type in ["figure", "table"]:
                layout_blocks_filtered.append(block)
            else:
                is_contained = False
                for img_block in image_blocks + table_blocks:
                    if is_bbox_contained(block.bbox, img_block.bbox):
                        is_contained = True
                        break
                if not is_contained:
                    layout_blocks_filtered.append(block)
        layout_blocks = layout_blocks_filtered
        
        # Remove empty text blocks
        layout_blocks_cleaned = []
        for block in layout_blocks:
            if block.block_type in ["figure", "table"]:
                layout_blocks_cleaned.append(block)
            elif hasattr(block, "elements") and len(block.elements) > 0:
                layout_blocks_cleaned.append(block)
        layout_blocks = layout_blocks_cleaned
        
        # Sort blocks: top-to-bottom, then left-to-right
        layout_blocks = sorted(layout_blocks, key=lambda b: (b.bbox[1], b.bbox[0]))

        # Compute the last content y1 on this page (exclude abandon blocks).
        # This is used as state for detecting intentional page breaks at the next page boundary.
        content_blocks_for_y1 = [b for b in layout_blocks if b.block_type != "abandon"]
        this_page_last_content_y1 = max((b.bbox[3] for b in content_blocks_for_y1), default=0.0)
        
        # Store page index and page height in each block
        for block in layout_blocks:
            if not hasattr(block, 'extra'):
                block.extra = {}
            block.extra['page_idx'] = page_idx
            block.extra['page_height'] = page_height
        
        # Group blocks into rows
        rows = []
        current_row = []
        row_threshold = 10  # points
        for b in layout_blocks:
            top_y = b.bbox[1]
            if not current_row:
                current_row = [b]
                current_top = top_y
            else:
                if abs(top_y - current_top) <= row_threshold:
                    current_row.append(b)
                else:
                    rows.append(sorted(current_row, key=lambda x: x.bbox[0]))
                    current_row = [b]
                    current_top = top_y
        if current_row:
            rows.append(sorted(current_row, key=lambda x: x.bbox[0]))
        
        print(f"  Grouped into {len(rows)} rows")
        
        # Detect and replicate intentional author page breaks for new top-level sections.
        # Rule: only when the first content block on the page starts a new top-level section
        # (not a subsection),
        # and the last content block on the previous page ended far from the bottom.
        first_content_block = next((b for b in layout_blocks if b.block_type != "abandon"), None)
        if (
            first_content_block is not None
            and prev_page_height is not None
            and prev_page_last_content_y1 is not None
        ):
            heading_text = getattr(first_content_block, "text", "") or ""
            # In this pipeline, many section headers are detected as "title" and only later
            # promoted to "section_header" inside process_text_block after the first title.
            starts_section = (
                first_content_block.block_type == "section_header"
                or (first_content_block.block_type == "title" and first_title_processed)
            )
            heading_level = get_section_heading_level(heading_text, default_level=1)
            is_top_level_section = starts_section and (heading_level == 1)

            prev_page_gap_to_bottom = float(prev_page_height) - float(prev_page_last_content_y1)
            prev_page_not_near_bottom = prev_page_gap_to_bottom >= self.intentional_page_break_gap_threshold_pt

            if is_top_level_section and prev_page_not_near_bottom:
                # Avoid creating a leading blank page break when the document is still empty.
                if len(docx_doc.paragraphs) > 0 or len(docx_doc.tables) > 0:
                    docx_doc.add_page_break()

        # Process rows
        section = docx_doc.sections[0]
        for row in rows:
            row.sort(key=lambda b: b.bbox[0])
            
            # Check if row should be a table
            use_table = False
            if len(row) >= 2:
                all_pairs_ok = True
                for i in range(len(row) - 1):
                    if not (is_same_line(row[i], row[i+1]) and horizontally_separated(row[i], row[i+1], min_gap=20)):
                        all_pairs_ok = False
                        break
                use_table = all_pairs_ok
            
            # Process row
            if use_table and len(row) >= 2:
                prev_row_y1 = process_table_row(docx_doc, row, section, page_width_pts)
                continue
            
            # Process each block in row
            for block in row:
                block_type = block.block_type
                
                if block_type == "abandon":
                    continue
                
                # Image / figure handling
                if block_type == "figure":
                    process_figure_block(
                        docx_doc, block, page_image, scale_x, scale_y,
                        self.max_image_width, page_idx
                    )
                    last_text_paragraph = None
                    prev_text_block = None
                    continue
                
                # Table handling
                if block_type == "table":
                    process_table_block(
                        docx_doc, block, page_image, scale_x, scale_y,
                        self.max_image_width, page_idx
                    )
                    last_text_paragraph = None
                    prev_text_block = None
                    continue
                
                # Check if we should merge with previous text block
                should_merge = False
                if prev_text_block is not None and last_text_paragraph is not None:
                    should_merge = should_merge_with_previous_block(prev_text_block, block)
                
                # Text handling
                result = process_text_block(
                    docx_doc, block, block_type, first_title_processed,
                    self.style_map, doc_left_margin_in, row, prev_row_y1,
                    last_text_paragraph if should_merge else None
                )
                
                if result[0] is None:  # Skip if style_info is None
                    continue
                
                p, block_type, first_title_processed, y1_pdf, should_merge = result
                
                # Update state
                last_text_paragraph = p
                prev_text_block = block
                prev_row_y1 = max(prev_row_y1, y1_pdf)
            
            # Update prev_row_y1 to the maximum y1 of all blocks in the row
            prev_row_y1 = max(b.bbox[3] for b in row)
        
        return {
            "first_title_processed": first_title_processed,
            "prev_row_y1": prev_row_y1,
            "last_text_paragraph": last_text_paragraph,
            "prev_text_block": prev_text_block,
            "prev_page_height": page_height,
            "prev_page_last_content_y1": this_page_last_content_y1,
        }
