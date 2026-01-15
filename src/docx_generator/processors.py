"""
Processors for different block types in DOCX generation.
"""

import os
from typing import Dict, List, Optional, Tuple

import cv2
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from typing import Dict, List, Optional, Tuple

import cv2
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

try:
    from src.utils import (
        clean_font_name,
        round_font_size,
        get_section_heading_level,
        remove_table_borders,
        set_table_col_widths,
        is_same_line,
        horizontally_separated,
    )
    from src.yolo.iou_matching import TextElement, LayoutBlock
except ImportError:
    # Fallback for direct execution
    import sys
    from pathlib import Path
    sys.path.insert(0, str(Path(__file__).parent.parent))
    from utils import (
        clean_font_name,
        round_font_size,
        get_section_heading_level,
        remove_table_borders,
        set_table_col_widths,
        is_same_line,
        horizontally_separated,
    )
    from yolo.iou_matching import TextElement, LayoutBlock


def should_merge_with_previous_block(prev_block, curr_block):
    """
    Check if current text block should be merged with previous block into same paragraph.
    
    Args:
        prev_block: Previous LayoutBlock object
        curr_block: Current LayoutBlock object
        
    Returns:
        bool: True if blocks should be merged
    """
    if not prev_block or not curr_block:
        return False
    
    # Get text from blocks (preserve original text with spaces)
    prev_text = prev_block.text if hasattr(prev_block, 'text') else ""
    curr_text = curr_block.text if hasattr(curr_block, 'text') else ""
    
    if not prev_text or not curr_text:
        return False
    
    # Check if previous block ends with space (indicating continuation)
    # This is the primary indicator that blocks should be merged
    prev_ends_with_space = prev_text.rstrip().endswith(" ")
    
    # Also check if previous block doesn't end with sentence-ending punctuation
    prev_no_sentence_end = prev_text.rstrip()[-1] not in ".!?"
    
    # Merge if: previous ends with space OR (previous doesn't end sentence AND current doesn't start with capital)
    # This handles cases like "Mô hình " + "YOLO" (merge) and "Sentence." + "New" (don't merge)
    if prev_ends_with_space:
        return True  # Always merge if previous ends with space
    
    # If previous doesn't end sentence and current doesn't start with capital, likely continuation
    curr_starts_lowercase = curr_text.strip() and curr_text.strip()[0].islower()
    return prev_no_sentence_end and curr_starts_lowercase


def process_figure_block(docx_doc, block, page_image, scale_x, scale_y, max_image_width, page_idx):
    """
    Process a figure block by extracting and inserting the image.
    
    Args:
        docx_doc: Document object
        block: LayoutBlock object with block_type == "figure"
        page_image: NumPy array of the page image (RGB) - used as fallback
        scale_x, scale_y: Scale factors from PDF to image coordinates - used as fallback
        max_image_width: Maximum width for the image (Inches)
        page_idx: Page index for unique temp file name
        
    Returns:
        bool: True if image was successfully added, False otherwise
    """
    # Check if image path is available in extra (from PDF extraction)
    image_path = block.extra.get("image_path") if hasattr(block, "extra") else None
    
    if image_path and os.path.exists(image_path):
        # Use pre-extracted image from PDF
        p = docx_doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=max_image_width)
        return True
    
    # Fallback: Crop image from page_image (for cases where image_path is not available)
    if page_image is None:
        return False
    
    x0_pdf, y0_pdf, x1_pdf, y1_pdf = block.bbox
    x0 = int(max(0, x0_pdf * scale_x))
    y0 = int(max(0, y0_pdf * scale_y))
    x1 = int(min(page_image.shape[1], x1_pdf * scale_x))
    y1 = int(min(page_image.shape[0], y1_pdf * scale_y))
    
    crop_img = page_image[y0:y1, x0:x1]
    if crop_img.size == 0:
        return False
    
    temp_image_path = f"temp_crop_page{page_idx}.png"
    cv2.imwrite(temp_image_path, cv2.cvtColor(crop_img, cv2.COLOR_RGB2BGR))
    p = docx_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(temp_image_path, width=max_image_width)
    
    # Clean up temp file
    if os.path.exists(temp_image_path):
        os.remove(temp_image_path)
    
    return True


def process_table_block(docx_doc, block, page_image, scale_x, scale_y, max_image_width, page_idx):
    """
    Process a table block by cropping and inserting the table as an image.
    
    Args:
        docx_doc: Document object
        block: LayoutBlock object with block_type == "table"
        page_image: NumPy array of the page image (RGB)
        scale_x, scale_y: Scale factors from PDF to image coordinates
        max_image_width: Maximum width for the image (Inches)
        page_idx: Page index for unique temp file name
        
    Returns:
        bool: True if table image was successfully added, False otherwise
    """
    if page_image is None:
        return False
    
    # Crop table from page_image using bbox from YOLO
    x0_pdf, y0_pdf, x1_pdf, y1_pdf = block.bbox
    x0 = int(max(0, x0_pdf * scale_x))
    y0 = int(max(0, y0_pdf * scale_y))
    x1 = int(min(page_image.shape[1], x1_pdf * scale_x))
    y1 = int(min(page_image.shape[0], y1_pdf * scale_y))
    
    crop_img = page_image[y0:y1, x0:x1]
    if crop_img.size == 0:
        return False
    
    # Save cropped table as temporary image
    temp_image_path = f"temp_table_page{page_idx}.png"
    cv2.imwrite(temp_image_path, cv2.cvtColor(crop_img, cv2.COLOR_RGB2BGR))
    
    # Insert table image into DOCX
    p = docx_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(temp_image_path, width=max_image_width)
    
    # Clean up temp file
    if os.path.exists(temp_image_path):
        os.remove(temp_image_path)
    
    return True


def process_table_row(docx_doc, row, section, page_width_pts):
    """
    Process a row of blocks as a table.
    
    Args:
        docx_doc: Document object
        row: List of LayoutBlock objects to be placed in table
        section: Document section for margin info
        page_width_pts: Page width in points
        
    Returns:
        float: y1 coordinate of the row (for spacing calculation)
    """
    min_x_pt = min(b.bbox[0] for b in row)
    max_x_pt = max(b.bbox[2] for b in row)
    doc_left_margin_pt = section.left_margin.pt
    doc_right_margin_pt = section.right_margin.pt
    
    table_start_x = doc_left_margin_pt
    table_end_x = page_width_pts - doc_right_margin_pt
    table_total_width_pt = table_end_x - table_start_x
    
    num_cols = len(row)
    table = docx_doc.add_table(rows=1, cols=num_cols)
    table.autofit = False
    remove_table_borders(table)
    
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.append(tblPr)
    
    for e in tblPr.findall(qn("w:tblInd")):
        tblPr.remove(e)
    
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), "0")
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)
    
    # Calculate column widths
    block_widths = []
    block_positions = []
    for b in row:
        block_w = max(1.0, (b.bbox[2] - b.bbox[0]))
        block_widths.append(block_w)
        block_positions.append((b.bbox[0], b.bbox[2]))
    
    first_block_start = block_positions[0][0]
    last_block_end = block_positions[-1][1]
    total_block_span = last_block_end - first_block_start
    
    col_widths_pt = []
    if total_block_span > 0 and len(block_widths) > 0:
        for i, block_w in enumerate(block_widths):
            proportion = block_w / total_block_span
            col_width = proportion * table_total_width_pt
            col_widths_pt.append(max(1.0, col_width))
        
        current_total = sum(col_widths_pt)
        if current_total < table_total_width_pt:
            remaining = table_total_width_pt - current_total
            for i in range(len(col_widths_pt)):
                proportion = block_widths[i] / sum(block_widths) if sum(block_widths) > 0 else 1.0 / len(block_widths)
                col_widths_pt[i] += proportion * remaining
    else:
        col_width_per_col = table_total_width_pt / num_cols
        col_widths_pt = [col_width_per_col] * num_cols
    
    set_table_col_widths(table, col_widths_pt)
    
    # Fill table cells with full text formatting
    for col_idx, block in enumerate(row):
        cell = table.rows[0].cells[col_idx]
        cell.text = ""
        
        if not block.elements:
            continue
        
        # Group elements by line (same y coordinate)
        elems = block.elements
        spans = [{"bbox": e.bbox, "text": e.text, "element": e} for e in elems]
        spans = sorted(spans, key=lambda s: (s["bbox"][1], s["bbox"][0]))
        
        lines = []
        current_line = [spans[0]]
        y_tolerance_pt = 2.0
        
        for s in spans[1:]:
            prev = current_line[-1]
            y_prev = prev["bbox"][1]
            y_curr = s["bbox"][1]
            if abs(y_curr - y_prev) > y_tolerance_pt:
                lines.append(current_line)
                current_line = [s]
            else:
                current_line.append(s)
        lines.append(current_line)
        
        # Process each line as a paragraph with full formatting
        for line in lines:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add each element as a run with its own formatting
            for span in line:
                elem = span["element"]
                text_content = span["text"]
                if not text_content.endswith(" "):
                    text_content += " "
                
                run = p.add_run(text_content)
                
                # Apply font name
                run.font.name = "Times New Roman"
                
                # Apply font size (rounded to nearest 0.5)
                if elem.font_size is not None:
                    rounded_size = round_font_size(elem.font_size)
                    run.font.size = Pt(rounded_size)
                
                # Apply font flags (bold, italic)
                if elem.font_flags is not None:
                    run.bold = (elem.font_flags & 16) != 0
                    run.italic = (elem.font_flags & 8) != 0
    
    return max(b.bbox[3] for b in row)


def process_text_block(
    docx_doc, block, block_type, first_title_processed, 
    style_map, doc_left_margin_in, row, prev_row_y1, last_paragraph=None
):
    """
    Process a text block (title, plain text, captions, etc.).
    
    Args:
        docx_doc: Document object
        block: LayoutBlock object
        block_type: String block type (may be modified for title->section_header)
        first_title_processed: Boolean flag (will be modified if title is first)
        style_map: Dictionary mapping block types to styles
        doc_left_margin_in: Left margin in inches
        row: List of blocks in current row (for spacing calculation)
        prev_row_y1: Previous row's y1 coordinate
        last_paragraph: Previous paragraph object (for merging consecutive blocks)
        
    Returns:
        tuple: (paragraph object, modified block_type, modified first_title_processed, y1, should_merge)
    """
    # Handle title: first title is title, subsequent titles are headings
    if block_type == "title":
        if not first_title_processed:
            first_title_processed = True
        else:
            block_type = "section_header"
    
    style_info = style_map.get(block_type, {"style": "Normal"})
    if style_info is None:
        return None, block_type, first_title_processed, block.bbox[3], False
    
    # Extract bbox coordinates early (needed for return value)
    x0_pdf, y0_pdf, x1_pdf, y1_pdf = block.bbox
    
    # Get section info for right margin and page width (calculate once for reuse)
    section = docx_doc.sections[0]
    try:
        doc_right_margin_in = section.right_margin.inches
    except Exception:
        doc_right_margin_in = section.right_margin.pt / 72.0
    page_width_in = section.page_width.pt / 72.0
    
    # Calculate indents based on PDF coordinates
    # Reduce by 1 point (1/72 inches) for both left and right indents
    x0_in = x0_pdf / 72
    x1_in = x1_pdf / 72
    indent_from_margin_in = max(0.0, x0_in - doc_left_margin_in - (5.0 / 72.0))
    right_edge_in = page_width_in - doc_right_margin_in
    right_indent_from_margin_in = max(0.0, right_edge_in - x1_in - (5.0 / 72.0))
    
    # Check if we should merge with previous paragraph
    # should_merge is determined by caller and passed via last_paragraph parameter
    # If last_paragraph is provided, we merge; otherwise create new paragraph
    should_merge = (last_paragraph is not None)
    
    style_name = style_info["style"]
    
    # Use last_paragraph if merging, otherwise create new paragraph
    if should_merge and last_paragraph is not None:
        p = last_paragraph
        # Don't modify formatting when merging
    else:
        if style_name.startswith("Heading"):
            # Base level from style name (e.g., "Heading 1")
            base_level = int(style_name.split()[-1])
            # For section headers, refine level based on numbering pattern
            if block_type == "section_header":
                heading_text = getattr(block, "text", "") or ""
                heading_level = get_section_heading_level(heading_text, default_level=base_level)
            else:
                heading_level = base_level
            p = docx_doc.add_heading(level=heading_level)
        else:
            p = docx_doc.add_paragraph(style=style_name)
        
        # Positioning based on PDF coordinates
        # Left indent: distance from left margin to block's left edge
        p.paragraph_format.left_indent = Inches(indent_from_margin_in)
        
        # Right indent: distance from block's right edge to right margin
        p.paragraph_format.right_indent = Inches(right_indent_from_margin_in)
        
        # Set space_before for the 1st block in a row
        if block == row[0] and prev_row_y1 > 0:
            vertical_gap = max(0, (y0_pdf - prev_row_y1))
            p.paragraph_format.space_before = Pt(vertical_gap)
        else:
            p.paragraph_format.space_before = Pt(0)
        
        p.paragraph_format.space_after = Pt(2)
        
        # Set paragraph alignment
        if block_type == "title":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif block_type in ["figure_caption", "table_caption"]:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif block_type not in ["formula_caption"] and not style_name.startswith("Heading"):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add text runs with formatting
    # Group elements into text blocks using actual text block information from PyMuPDF
    if not block.elements:
        return p, block_type, first_title_processed, y1_pdf, should_merge
    
    # Check if we have text block information in extra
    span_to_block_idx = block.extra.get("span_to_block_idx")
    
    if span_to_block_idx is not None:
        # Use actual text block information from PyMuPDF
        # Group elements by their text block index, preserving original order within each block
        text_block_groups: Dict[int, List[Tuple[TextElement, int]]] = {}  # block_idx -> list of (element, original_span_idx)
        
        span_to_original_order = block.extra.get("span_to_original_order")
        
        for elem_idx, elem in enumerate(block.elements):
            block_idx = span_to_block_idx.get(elem_idx, 0)  # Default to 0 if not found
            if block_idx not in text_block_groups:
                text_block_groups[block_idx] = []
            
            # Get original span index within the block
            if span_to_original_order and elem_idx in span_to_original_order:
                _, original_span_idx = span_to_original_order[elem_idx]
            else:
                # Fallback: use bbox-based ordering
                original_span_idx = elem_idx
            
            text_block_groups[block_idx].append((elem, original_span_idx))
        
        # Sort by block index to maintain order
        sorted_block_indices = sorted(text_block_groups.keys())
        
        # Process each text block group as a separate paragraph
        for group_idx, block_idx in enumerate(sorted_block_indices):
            # Sort elements within each group by original span order within the text block
            # This preserves the correct order of spans as they appear in the original text block
            group_with_order = sorted(text_block_groups[block_idx], key=lambda x: x[1])
            group = [elem for elem, _ in group_with_order]
            
            if group_idx > 0:
                # Create new paragraph for each text block after the first
                p = docx_doc.add_paragraph(style=style_name)
                # Copy formatting from first paragraph
                p.paragraph_format.left_indent = Inches(indent_from_margin_in)
                p.paragraph_format.right_indent = Inches(right_indent_from_margin_in)
                p.paragraph_format.space_after = Pt(2)
                if block_type == "title":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif block_type in ["figure_caption", "table_caption"]:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif block_type not in ["formula_caption"] and not style_name.startswith("Heading"):
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Add elements in this group
            for elem_idx, elem in enumerate(group):
                text_content = elem.text
                if not text_content.endswith(" "):
                    text_content += " "
                run = p.add_run(text_content)
                # Set all text to Times New Roman for testing
                run.font.name = "Times New Roman"
                if elem.font_size is not None:
                    rounded_size = round_font_size(elem.font_size)
                    run.font.size = Pt(rounded_size)
                if elem.font_flags is not None:
                    run.bold = (elem.font_flags & 16) != 0
                    run.italic = (elem.font_flags & 8) != 0
    else:
        # Fallback: use vertical spacing if text block info is not available
        # Calculate average font size for threshold
        font_sizes = [e.font_size for e in block.elements if e.font_size is not None]
        avg_font_size = sum(font_sizes) / len(font_sizes) if font_sizes else 12.0
        vertical_gap_threshold = avg_font_size * 1.5  # Threshold for detecting new text block
        
        # Group elements into text blocks
        text_block_groups = []
        current_group = [block.elements[0]]
        
        for i in range(1, len(block.elements)):
            prev_elem = block.elements[i - 1]
            curr_elem = block.elements[i]
            
            # Calculate vertical gap (y0 of current - y1 of previous)
            vertical_gap = curr_elem.bbox[1] - prev_elem.bbox[3]
            
            if vertical_gap > vertical_gap_threshold:
                # Large gap detected - start new text block
                text_block_groups.append(current_group)
                current_group = [curr_elem]
            else:
                # Same text block - add to current group
                current_group.append(curr_elem)
        
        # Add last group
        if current_group:
            text_block_groups.append(current_group)
        
        # Process each text block group as a separate paragraph
        for group_idx, group in enumerate(text_block_groups):
            if group_idx > 0:
                # Create new paragraph for each text block after the first
                p = docx_doc.add_paragraph(style=style_name)
                # Copy formatting from first paragraph
                p.paragraph_format.left_indent = Inches(indent_from_margin_in)
                p.paragraph_format.right_indent = Inches(right_indent_from_margin_in)
                p.paragraph_format.space_after = Pt(2)
                if block_type == "title":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif block_type in ["figure_caption", "table_caption"]:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif block_type not in ["formula_caption"] and not style_name.startswith("Heading"):
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Add elements in this group
            for elem_idx, elem in enumerate(group):
                text_content = elem.text
                if not text_content.endswith(" "):
                    text_content += " "
                run = p.add_run(text_content)
                # Set all text to Times New Roman for testing
                run.font.name = "Times New Roman"
                if elem.font_size is not None:
                    rounded_size = round_font_size(elem.font_size)
                    run.font.size = Pt(rounded_size)
                if elem.font_flags is not None:
                    run.bold = (elem.font_flags & 16) != 0
                    run.italic = (elem.font_flags & 8) != 0
    
    # Return the last paragraph created (p) and merge flag
    return p, block_type, first_title_processed, y1_pdf, should_merge
